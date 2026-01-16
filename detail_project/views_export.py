# detail_project/views_export.py
"""
Export views for Jadwal Pekerjaan
Handles batch export API untuk PDF dan Word generation
"""
from django.http import JsonResponse, FileResponse, Http404
from django.contrib.auth.decorators import login_required
from django.views.decorators.http import require_http_methods
from django.views.decorators.csrf import csrf_exempt
from django.shortcuts import get_object_or_404
from django.utils import timezone
from django.core.files.base import ContentFile
from django.db import transaction

import json
import base64
import logging

from .models_export import ExportSession, ExportPage

logger = logging.getLogger(__name__)


# ============================================================================
# API Endpoints for Batch Export (PDF/Word)
# ============================================================================

@login_required
@csrf_exempt  # Frontend will send CSRF token via header
@require_http_methods(["POST"])
def export_init(request):
    """
    Initialize export session

    POST /api/export/init
    Body:
    {
        "reportType": "rekap" | "monthly" | "weekly",
        "format": "pdf" | "word",
        "estimatedPages": 100,
        "projectName": "Nama Project",
        "metadata": {
            "month": 2,
            "week": 5,
            "options": {...}
        }
    }

    Returns:
    {
        "exportId": "abc-123-def",
        "expiresAt": "2025-12-17T12:00:00Z"
    }
    """
    try:
        # Parse request body
        data = json.loads(request.body)

        report_type = data.get('reportType')
        format_type = data.get('format')
        estimated_pages = data.get('estimatedPages', 0)
        project_name = data.get('projectName', '')
        metadata = data.get('metadata', {})

        # Validate required fields
        if not report_type:
            return JsonResponse({'error': 'reportType is required'}, status=400)

        if not format_type:
            return JsonResponse({'error': 'format is required'}, status=400)

        # Validate report type
        valid_report_types = [choice[0] for choice in ExportSession.REPORT_CHOICES]
        if report_type not in valid_report_types:
            return JsonResponse({
                'error': f'Invalid reportType. Must be one of: {", ".join(valid_report_types)}'
            }, status=400)

        # Validate format
        valid_formats = [choice[0] for choice in ExportSession.FORMAT_CHOICES]
        if format_type not in valid_formats:
            return JsonResponse({
                'error': f'Invalid format. Must be one of: {", ".join(valid_formats)}'
            }, status=400)

        # ============================================================
        # SUBSCRIPTION CHECK - Export Restrictions
        # ============================================================
        # Feature Matrix:
        # - TRIAL: ❌ (cannot export at all)
        # - PRO: ✅ (clean exports)
        # - EXPIRED: PDF with watermark only, no Word/Excel
        
        user = request.user
        add_watermark = False
        
        if user.subscription_status == 'TRIAL':
            return JsonResponse({
                'success': False,
                'error': 'Fitur export tersedia setelah Anda upgrade ke Pro. Trial tidak termasuk fitur export.',
                'code': 'TRIAL_NO_EXPORT',
                'upgrade_url': '/subscriptions/pricing/'
            }, status=403)
        
        if not user.is_pro_active:
            # EXPIRED user
            if format_type in ['word', 'xlsx']:
                return JsonResponse({
                    'success': False,
                    'error': 'Export Word dan Excel hanya tersedia untuk pengguna Pro. Silakan upgrade langganan Anda.',
                    'code': 'PRO_REQUIRED',
                    'subscription_status': user.subscription_status,
                    'upgrade_url': '/subscriptions/pricing/'
                }, status=403)
            # PDF allowed but with watermark
            add_watermark = True

        # Add watermark flag to metadata for PDF generation
        if add_watermark:
            metadata['add_watermark'] = True
            metadata['watermark_text'] = 'DEMO - Dashboard-RAB'

        # Create export session
        session = ExportSession.objects.create(
            user=request.user,
            report_type=report_type,
            format_type=format_type,
            estimated_pages=estimated_pages,
            project_name=project_name,
            metadata=metadata
        )

        logger.info(f"Export session created: {session.export_id} for user {request.user.id}")

        return JsonResponse({
            'exportId': str(session.export_id),
            'expiresAt': session.expires_at.isoformat()
        }, status=201)

    except json.JSONDecodeError:
        return JsonResponse({'error': 'Invalid JSON'}, status=400)

    except Exception as e:
        logger.error(f"Error in export_init: {str(e)}", exc_info=True)
        return JsonResponse({'error': 'Internal server error'}, status=500)


@login_required
@csrf_exempt
@require_http_methods(["POST"])
def export_upload_pages(request):
    """
    Upload batch of pages to export session

    POST /api/export/upload-pages
    Body:
    {
        "exportId": "abc-123-def",
        "batchIndex": 0,
        "pages": [
            {
                "pageNumber": 1,
                "title": "Gantt W1-W12",
                "dataUrl": "data:image/png;base64,...",
                "format": "png"
            },
            ...
        ]
    }

    Returns:
    {
        "received": 10,
        "totalReceived": 10,
        "progress": 10.0
    }
    """
    try:
        # Parse request body
        data = json.loads(request.body)

        export_id = data.get('exportId')
        batch_index = data.get('batchIndex', 0)
        pages = data.get('pages', [])

        # Validate required fields
        if not export_id:
            return JsonResponse({'error': 'exportId is required'}, status=400)

        if not pages:
            return JsonResponse({'error': 'pages array is required and cannot be empty'}, status=400)

        # Get export session (must belong to current user)
        session = get_object_or_404(
            ExportSession,
            export_id=export_id,
            user=request.user
        )

        # Check if session expired
        if session.is_expired:
            session.mark_failed('Session expired')
            return JsonResponse({'error': 'Export session has expired'}, status=410)  # 410 Gone

        # Check if session already completed or failed
        if session.status in [ExportSession.STATUS_COMPLETED, ExportSession.STATUS_FAILED]:
            return JsonResponse({
                'error': f'Session already {session.status}'
            }, status=400)

        # Mark as uploading (if first batch)
        if session.status == ExportSession.STATUS_INIT:
            session.mark_uploading()

        # Save pages to database
        pages_saved = 0
        with transaction.atomic():
            for page_data in pages:
                page_number = page_data.get('pageNumber')
                title = page_data.get('title', '')
                data_url = page_data.get('dataUrl', '')
                format_type = page_data.get('format', 'png')

                if not page_number or not data_url:
                    logger.warning(f"Skipping invalid page in batch {batch_index}: missing pageNumber or dataUrl")
                    continue

                # Create or update page
                ExportPage.objects.update_or_create(
                    session=session,
                    page_number=page_number,
                    defaults={
                        'batch_index': batch_index,
                        'title': title,
                        'image_data': data_url,
                        'format': format_type,
                        'file_size': len(data_url)
                    }
                )

                pages_saved += 1

            # Update session counters
            session.increment_pages(pages_saved)
            session.increment_batches()

        logger.info(f"Uploaded {pages_saved} pages for session {export_id} (batch {batch_index})")

        return JsonResponse({
            'received': pages_saved,
            'totalReceived': session.pages_received,
            'progress': session.progress_percent
        })

    except json.JSONDecodeError:
        return JsonResponse({'error': 'Invalid JSON'}, status=400)

    except Exception as e:
        logger.error(f"Error in export_upload_pages: {str(e)}", exc_info=True)
        return JsonResponse({'error': 'Internal server error'}, status=500)


@login_required
@csrf_exempt
@require_http_methods(["POST"])
def export_finalize(request):
    """
    Finalize export and generate PDF/Word file

    POST /api/export/finalize
    Body:
    {
        "exportId": "abc-123-def"
    }

    Returns:
    {
        "downloadUrl": "/media/exports/2025/12/16/abc-123-def.pdf",
        "fileSize": 5242880,
        "pagesGenerated": 100
    }
    """
    try:
        # Parse request body
        data = json.loads(request.body)

        export_id = data.get('exportId')

        # Validate required fields
        if not export_id:
            return JsonResponse({'error': 'exportId is required'}, status=400)

        # Get export session (must belong to current user)
        session = get_object_or_404(
            ExportSession,
            export_id=export_id,
            user=request.user
        )

        # Check if session expired
        if session.is_expired:
            session.mark_failed('Session expired')
            return JsonResponse({'error': 'Export session has expired'}, status=410)

        # Check if already completed
        if session.status == ExportSession.STATUS_COMPLETED:
            return JsonResponse({
                'downloadUrl': session.download_url,
                'fileSize': session.file_size,
                'pagesGenerated': session.pages_received
            })

        # Check if already failed
        if session.status == ExportSession.STATUS_FAILED:
            return JsonResponse({
                'error': f'Export failed: {session.error_message}'
            }, status=400)

        # Mark as processing
        session.mark_processing()

        # Get all pages ordered by page number
        pages = session.pages.all().order_by('page_number')

        if not pages.exists():
            session.mark_failed('No pages uploaded')
            return JsonResponse({'error': 'No pages uploaded'}, status=400)

        # Generate PDF, Word in Excel based on format
        try:
            if session.format_type == ExportSession.FORMAT_PDF:
                output_file, file_size = generate_pdf_from_pages(session, pages)
            elif session.format_type == ExportSession.FORMAT_WORD:
                output_file, file_size = generate_word_from_pages(session, pages)
            elif session.format_type == ExportSession.FORMAT_XLSX:
                output_file, file_size = generate_excel_from_pages(session, pages)
            else:
                raise ValueError(f"Unsupported format: {session.format_type}")

            # Mark as completed
            session.mark_completed(output_file, file_size)

            logger.info(f"Export session {export_id} completed successfully")

            return JsonResponse({
                'downloadUrl': session.download_url,
                'fileSize': session.file_size,
                'pagesGenerated': session.pages_received
            })

        except Exception as e:
            error_msg = str(e)
            logger.error(f"Error generating {session.format_type} for session {export_id}: {error_msg}", exc_info=True)
            session.mark_failed(error_msg)
            return JsonResponse({'error': f'Generation failed: {error_msg}'}, status=500)

    except json.JSONDecodeError:
        return JsonResponse({'error': 'Invalid JSON'}, status=400)

    except Exception as e:
        logger.error(f"Error in export_finalize: {str(e)}", exc_info=True)
        return JsonResponse({'error': 'Internal server error'}, status=500)


# ============================================================================
# Helper Functions - PDF/Word Generation
# ============================================================================

def generate_pdf_from_pages(session, pages):
    """
    Generate PDF from PNG pages using ReportLab

    Args:
        session: ExportSession instance
        pages: QuerySet of ExportPage instances

    Returns:
        tuple: (ContentFile, file_size)
    """
    from io import BytesIO
    from reportlab.lib.pagesizes import A4, landscape
    from reportlab.pdfgen import canvas
    from reportlab.lib.utils import ImageReader
    from reportlab.lib.colors import Color

    # Check if watermark should be added (for non-PRO users)
    add_watermark = False
    watermark_text = 'DEMO - Dashboard-RAB'
    if session.metadata:
        add_watermark = session.metadata.get('add_watermark', False)
        watermark_text = session.metadata.get('watermark_text', watermark_text)

    # Create PDF buffer
    buffer = BytesIO()

    # Create PDF canvas (A4 landscape)
    pdf = canvas.Canvas(buffer, pagesize=landscape(A4))
    page_width, page_height = landscape(A4)

    # Add each page as image
    for page in pages:
        # Decode base64 image
        image_data = page.image_data
        if image_data.startswith('data:image/png;base64,'):
            image_data = image_data.split(',', 1)[1]

        image_bytes = base64.b64decode(image_data)
        image_buffer = BytesIO(image_bytes)

        # Add image to PDF (fit to page)
        try:
            img = ImageReader(image_buffer)
            pdf.drawImage(img, 0, 0, width=page_width, height=page_height, preserveAspectRatio=True, anchor='c')
            
            # Add watermark overlay if required
            if add_watermark:
                pdf.saveState()
                # Set watermark style - semi-transparent gray
                pdf.setFillColor(Color(0.5, 0.5, 0.5, alpha=0.3))
                pdf.setFont("Helvetica-Bold", 60)
                # Rotate and center the watermark diagonally
                pdf.translate(page_width / 2, page_height / 2)
                pdf.rotate(45)
                # Draw watermark text centered
                pdf.drawCentredString(0, 0, watermark_text)
                pdf.restoreState()
            
            pdf.showPage()  # Next page
        except Exception as e:
            logger.error(f"Error adding page {page.page_number} to PDF: {str(e)}")
            raise

    # Save PDF
    pdf.save()

    # Get PDF bytes
    pdf_bytes = buffer.getvalue()
    buffer.close()

    # Create filename
    filename = f"{session.export_id}.pdf"

    # Create ContentFile
    content_file = ContentFile(pdf_bytes, name=filename)
    file_size = len(pdf_bytes)

    return content_file, file_size


def generate_word_from_pages(session, pages):
    """
    Generate Word document from PNG pages using python-docx

    Args:
        session: ExportSession instance
        pages: QuerySet of ExportPage instances

    Returns:
        tuple: (ContentFile, file_size)
    """
    from io import BytesIO
    from docx import Document
    from docx.shared import Inches

    # Create Word document
    doc = Document()

    # Add each page as image
    for page in pages:
        # Decode base64 image
        image_data = page.image_data
        if image_data.startswith('data:image/png;base64,'):
            image_data = image_data.split(',', 1)[1]

        image_bytes = base64.b64decode(image_data)
        image_buffer = BytesIO(image_bytes)

        # Add image to document
        try:
            # Add page title if available
            if page.title:
                doc.add_heading(page.title, level=2)

            # Add image (fit to page width: 6.5 inches for letter size)
            doc.add_picture(image_buffer, width=Inches(6.5))

            # Page break after each image (except last)
            if page != pages.last():
                doc.add_page_break()

        except Exception as e:
            logger.error(f"Error adding page {page.page_number} to Word: {str(e)}")
            raise

    # Save to buffer
    buffer = BytesIO()
    doc.save(buffer)

    # Get Word bytes
    word_bytes = buffer.getvalue()
    buffer.close()

    # Create filename
    filename = f"{session.export_id}.docx"

    # Create ContentFile
    content_file = ContentFile(word_bytes, name=filename)
    file_size = len(word_bytes)

    return content_file, file_size


def generate_excel_from_pages(session, pages):
    """
    Generate Excel workbook using ExportManager for native charts.
    
    For Rekap reports, uses professional export with:
    - Cover sheet with project info
    - Kurva S sheet with data table and native LineChart
    - Input Progress-Gantt sheet (SSOT for input values)

    Args:
        session: ExportSession instance
        pages: QuerySet of ExportPage instances

    Returns:
        tuple: (ContentFile, file_size)
    """
    from io import BytesIO
    from django.core.files.base import ContentFile
    
    logger.info(f"[generate_excel_from_pages] session={session.export_id}, report_type={session.report_type}, metadata={session.metadata}")
    
    # Try to use ExportManager for professional export with native charts
    if session.report_type in ('rekap', 'monthly', 'weekly'):
        try:
            # Get project from session metadata
            project_id = session.metadata.get('projectId') if session.metadata else None
            logger.info(f"[generate_excel_from_pages] projectId from metadata: {project_id}")
            
            if project_id:
                from .models import Project
                project = Project.objects.get(id=project_id)
                logger.info(f"[generate_excel_from_pages] Found project: {project.nama}")
                
                from .exports.export_manager import ExportManager
                manager = ExportManager(project)
                
                # Use professional export for native chart support
                logger.info(f"[generate_excel_from_pages] Calling export_jadwal_professional...")
                response = manager.export_jadwal_professional(
                    format_type='xlsx',
                    report_type=session.report_type
                )
                
                # Extract bytes from response
                excel_bytes = response.content
                logger.info(f"[generate_excel_from_pages] Professional export success, size={len(excel_bytes)} bytes")
                
                # Create filename
                filename = f"{session.export_id}.xlsx"
                content_file = ContentFile(excel_bytes, name=filename)
                file_size = len(excel_bytes)
                
                logger.info(f"Generated professional Excel for session {session.export_id}")
                return content_file, file_size
            else:
                logger.warning(f"[generate_excel_from_pages] projectId not found in metadata")
                
        except Exception as e:
            logger.warning(f"Professional export failed, falling back to image-based: {str(e)}", exc_info=True)
            # Fall through to image-based export
    
    # Fallback: Simple image-based Excel (for non-rekap or if professional fails)
    from openpyxl import Workbook
    from openpyxl.drawing.image import Image as XLImage
    from openpyxl.styles import Font

    # Create Excel workbook
    wb = Workbook()
    ws = wb.active
    ws.title = "Ringkasan"

    # Add title
    ws['A1'] = f"Laporan {session.report_type.upper()} - {session.project_name}"
    ws['A1'].font = Font(size=16, bold=True)
    ws.merge_cells('A1:D1')

    # Add each page as image in separate sheet
    for page in pages:
        # Decode base64 image
        image_data = page.image_data
        if image_data.startswith('data:image/png;base64,'):
            image_data = image_data.split(',', 1)[1]

        image_bytes = base64.b64decode(image_data)
        image_buffer = BytesIO(image_bytes)

        # Create sheet for this image
        try:
            safe_title = (page.title or f"Page {page.page_number}")[:31]
            # Remove invalid characters
            safe_title = ''.join(c for c in safe_title if c not in '[]:*?/\\')
            ws_img = wb.create_sheet(title=safe_title)

            # Add image
            xl_img = XLImage(image_buffer)
            ws_img.add_image(xl_img, 'A1')

        except Exception as e:
            logger.error(f"Error adding page {page.page_number} to Excel: {str(e)}")
            # Continue with other pages
            continue

    # Save to buffer
    buffer = BytesIO()
    wb.save(buffer)

    # Get Excel bytes
    excel_bytes = buffer.getvalue()
    buffer.close()

    # Create filename
    filename = f"{session.export_id}.xlsx"

    # Create ContentFile
    content_file = ContentFile(excel_bytes, name=filename)
    file_size = len(excel_bytes)

    return content_file, file_size


# ============================================================================
# Download Endpoint
# ============================================================================

@login_required
@require_http_methods(["GET"])
def export_download(request, export_id):
    """
    Download completed export file

    GET /api/export/download/<export_id>

    Returns: File download (PDF or Word)
    """
    try:
        # Get export session (must belong to current user)
        session = get_object_or_404(
            ExportSession,
            export_id=export_id,
            user=request.user
        )

        # Check if completed
        if session.status != ExportSession.STATUS_COMPLETED:
            raise Http404("Export not completed")

        # Check if file exists
        if not session.output_file:
            raise Http404("Export file not found")

        # Return file response
        response = FileResponse(
            session.output_file.open('rb'),
            content_type='application/octet-stream'
        )

        # Set filename for download
        filename = f"{session.project_name or 'export'}_{session.report_type}.{session.format_type}"
        filename = filename.replace(' ', '_')  # Remove spaces
        response['Content-Disposition'] = f'attachment; filename="{filename}"'

        return response

    except Exception as e:
        logger.error(f"Error in export_download: {str(e)}", exc_info=True)
        raise Http404("Export file not found")


# ============================================================================
# Status Check Endpoint
# ============================================================================

@login_required
@require_http_methods(["GET"])
def export_status(request, export_id):
    """
    Get export session status

    GET /api/export/status/<export_id>

    Returns:
    {
        "status": "uploading" | "processing" | "completed" | "failed",
        "progress": 45.5,
        "pagesReceived": 45,
        "estimatedPages": 100,
        "downloadUrl": "/api/export/download/abc-123-def" (if completed)
    }
    """
    try:
        # Get export session (must belong to current user)
        session = get_object_or_404(
            ExportSession,
            export_id=export_id,
            user=request.user
        )

        response_data = {
            'status': session.status,
            'progress': session.progress_percent,
            'pagesReceived': session.pages_received,
            'estimatedPages': session.estimated_pages,
        }

        # Add download URL if completed
        if session.status == ExportSession.STATUS_COMPLETED:
            response_data['downloadUrl'] = f'/api/export/download/{export_id}'

        # Add error message if failed
        if session.status == ExportSession.STATUS_FAILED:
            response_data['error'] = session.error_message

        return JsonResponse(response_data)

    except Exception as e:
        logger.error(f"Error in export_status: {str(e)}", exc_info=True)
        return JsonResponse({'error': 'Internal server error'}, status=500)


# ============================================================================
# Async Export API Endpoints (Celery-based)
# ============================================================================

@login_required
@csrf_exempt
@require_http_methods(["POST"])
def api_start_export_async(request, project_id):
    """
    Start async export task for PDF/Word generation.
    
    POST /api/project/{project_id}/export-async/
    Body:
    {
        "export_type": "rekap-rab" | "jadwal-pekerjaan" | "harga-items" | 
                       "rincian-ahsp" | "rekap-kebutuhan" | "volume-pekerjaan",
        "format": "pdf" | "word",
        "options": {...}  // Optional export configuration
    }
    
    Returns:
    {
        "task_id": "abc-123-def",
        "status_url": "/api/export-status/async/abc-123-def/"
    }
    """
    from detail_project.tasks import generate_export_async
    from detail_project.models import Project
    
    try:
        # Validate project access
        project = get_object_or_404(Project, id=project_id)
        if project.owner != request.user:
            return JsonResponse({'error': 'Permission denied'}, status=403)
        
        # Parse request
        data = json.loads(request.body)
        export_type = data.get('export_type')
        format_type = data.get('format', 'pdf')
        options = data.get('options', {})
        
        # Validate export_type
        valid_types = [
            'rekap-rab', 'jadwal-pekerjaan', 'harga-items',
            'rincian-ahsp', 'rekap-kebutuhan', 'volume-pekerjaan'
        ]
        if export_type not in valid_types:
            return JsonResponse({
                'error': f'Invalid export_type. Must be one of: {", ".join(valid_types)}'
            }, status=400)
        
        # Validate format
        if format_type not in ['pdf', 'word']:
            return JsonResponse({'error': 'Invalid format. Must be pdf or word'}, status=400)
        
        # ============================================================
        # SUBSCRIPTION CHECK - Export Restrictions
        # ============================================================
        user = request.user
        add_watermark = False
        
        if user.subscription_status == 'TRIAL':
            return JsonResponse({
                'success': False,
                'error': 'Fitur export tersedia setelah Anda upgrade ke Pro. Trial tidak termasuk fitur export.',
                'code': 'TRIAL_NO_EXPORT',
                'upgrade_url': '/subscriptions/pricing/'
            }, status=403)
        
        if not user.is_pro_active:
            # EXPIRED user
            if format_type == 'word':
                return JsonResponse({
                    'success': False,
                    'error': 'Export Word hanya tersedia untuk pengguna Pro. Silakan upgrade langganan Anda.',
                    'code': 'PRO_REQUIRED',
                    'upgrade_url': '/subscriptions/pricing/'
                }, status=403)
            # PDF allowed but with watermark
            add_watermark = True
            options['add_watermark'] = True
            options['watermark_text'] = 'DEMO - Dashboard-RAB'
        
        # Start Celery task
        logger.info(
            f"Starting async export: project={project_id}, type={export_type}, "
            f"format={format_type}, user={request.user.id}"
        )
        
        task = generate_export_async.delay(
            project_id=project_id,
            export_type=export_type,
            format_type=format_type,
            user_id=request.user.id,
            options=options
        )
        
        return JsonResponse({
            'task_id': task.id,
            'status_url': f'/api/export-status/async/{task.id}/'
        }, status=202)  # 202 Accepted
    
    except json.JSONDecodeError:
        return JsonResponse({'error': 'Invalid JSON'}, status=400)
    
    except Exception as e:
        logger.error(f"Error starting async export: {str(e)}", exc_info=True)
        return JsonResponse({'error': 'Internal server error'}, status=500)


@login_required
@require_http_methods(["GET"])
def api_export_status_async(request, task_id):
    """
    Check status of async export task.
    
    GET /api/export-status/async/{task_id}/
    
    Returns:
    {
        "task_id": "abc-123-def",
        "status": "PENDING" | "PROCESSING" | "SUCCESS" | "FAILURE",
        "progress": {
            "status": "Generating rekap-rab PDF...",
            "progress": 45,
            "total": 100
        },  // Only if PROCESSING
        "result": {
            "file_path": "/path/to/export.pdf",
            "file_size": 1024000,
            "download_url": "/api/export-download/async/abc-123-def/"
        },  // Only if SUCCESS
        "error": "Error message"  // Only if FAILURE
    }
    """
    from celery.result import AsyncResult
    
    try:
        # Get task result
        task = AsyncResult(task_id)
        
        response_data = {
            'task_id': task_id,
            'status': task.state
        }
        
        # Add state-specific data
        if task.state == 'PENDING':
            response_data['message'] = 'Task is waiting to start...'
        
        elif task.state == 'PROCESSING':
            # Get progress info from task meta
            if task.info:
                response_data['progress'] = task.info
        
        elif task.state == 'SUCCESS':
            # Get result data
            result = task.result or {}
            response_data['result'] = {
                'export_type': result.get('export_type'),
                'format': result.get('format'),
                'file_size': result.get('file_size'),
                'download_url': f'/api/export-download/async/{task_id}/'
            }
        
        elif task.state == 'FAILURE':
            # Get error message
            error_info = task.info or {}
            if isinstance(error_info, dict):
                response_data['error'] = error_info.get('error', str(error_info))
            else:
                response_data['error'] = str(error_info)
        
        return JsonResponse(response_data)
    
    except Exception as e:
        logger.error(f"Error checking async export status: {str(e)}", exc_info=True)
        return JsonResponse({'error': 'Internal server error'}, status=500)


@login_required
@require_http_methods(["GET"])
def api_export_download_async(request, task_id):
    """
    Download completed async export file.
    
    GET /api/export-download/async/{task_id}/
    
    Returns: File download (PDF or Word)
    """
    from celery.result import AsyncResult
    import os
    
    try:
        # Get task result
        task = AsyncResult(task_id)
        
        # Check if task completed
        if task.state != 'SUCCESS':
            return JsonResponse({
                'error': f'Export not ready. Status: {task.state}'
            }, status=400)
        
        # Get file path from result
        result = task.result or {}
        file_path = result.get('file_path')
        
        if not file_path or not os.path.exists(file_path):
            raise Http404("Export file not found")
        
        # Determine content type
        format_type = result.get('format', 'pdf')
        content_type = 'application/pdf' if format_type == 'pdf' else 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        
        # Return file
        response = FileResponse(
            open(file_path, 'rb'),
            content_type=content_type
        )
        
        # Set filename
        export_type = result.get('export_type', 'export')
        filename = f"{export_type}.{format_type}"
        response['Content-Disposition'] = f'attachment; filename="{filename}"'
        
        logger.info(f"Async export downloaded: task_id={task_id}, file={filename}")
        
        return response
    
    except Exception as e:
        logger.error(f"Error downloading async export: {str(e)}", exc_info=True)
        raise Http404("Export file not found")