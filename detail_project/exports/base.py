# ============================================================================
# FILE: detail_project/exports/base.py
# ============================================================================
"""
Base Exporter - Shared Logic untuk Semua Export Formats

Menyediakan base class dan shared utilities untuk CSV, PDF, dan Word exports.
Semua specific exports (RekapRABExporter, RincianRABExporter, dll) 
akan inherit dari BaseExporter ini.

Architecture:
    BaseExporter (abstract)
    ├── CSV methods (shared)
    ├── PDF methods (shared)  
    ├── Word methods (shared)
    └── Style helpers (using export_config)

Author: Export Refactoring Phase 1
Created: 2025
"""

import csv
from io import BytesIO
from decimal import Decimal
from datetime import datetime
from django.http import HttpResponse
from django.utils.timezone import now

# Import from export_config (single source of truth)
from ..export_config import (
    ExportColors,
    ExportFonts,
    ExportLayout,
    get_level_style,
    format_currency,
    format_volume,
)

# Optional imports dengan fallback
try:
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import A4, landscape
    from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import mm
    REPORTLAB_AVAILABLE = True
except ImportError:
    REPORTLAB_AVAILABLE = False

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False


# ============================================================================
# BASE EXPORTER CLASS
# ============================================================================

class BaseExporter:
    """
    Base class untuk semua exports.
    
    Menyediakan shared methods untuk CSV, PDF, dan Word exports.
    Subclass harus implement specific logic di method to_csv(), to_pdf(), to_word().
    
    Attributes:
        project: Project instance
        data: Data rows (list of dicts)
        pricing: ProjectPricing instance
        timestamp: Timestamp saat export dibuat
    
    Usage:
        class RekapRABExporter(BaseExporter):
            def to_csv(self):
                response = self._create_csv_response('rekap.csv')
                writer = csv.writer(response, delimiter=';')
                # ... specific logic ...
                return response
    """
    
    def __init__(self, project, data, pricing):
        """
        Initialize exports dengan project data.
        
        Args:
            project: Project instance
            data: Data rows (list of dicts) - format depends on export type
            pricing: ProjectPricing instance (for markup, etc)
        """
        self.project = project
        self.data = data
        self.pricing = pricing
        self.timestamp = now()
    
    # ========================================================================
    # CSV SHARED METHODS
    # ========================================================================
    
    def _create_csv_response(self, base_filename: str) -> HttpResponse:
        """
        Create HTTP response untuk CSV export dengan UTF-8 BOM.
        
        Args:
            base_filename: Base filename tanpa timestamp (e.g., 'rekap_rab.csv')
        
        Returns:
            HttpResponse dengan proper headers untuk CSV
        
        Usage:
            response = self._create_csv_response('rekap_rab.csv')
            writer = csv.writer(response, delimiter=';')
        """
        response = HttpResponse(content_type='text/csv; charset=utf-8-sig')
        
        # Add timestamp to filename
        timestamp_str = self.timestamp.strftime('%Y%m%d_%H%M%S')
        filename = f"{base_filename.replace('.csv', '')}_{self.project.id}_{timestamp_str}.csv"
        
        response['Content-Disposition'] = f'attachment; filename="{filename}"'
        return response
    
    def _write_csv_header(self, writer, title: str, info_rows: list):
        """
        Write standard CSV header section.
        
        Args:
            writer: csv.writer instance
            title: Document title (e.g., 'REKAP RAB')
            info_rows: List of [label, ':', value] rows
        
        Usage:
            info_rows = [
                ['Nama Project', ':', self.project.nama],
                ['Profit/Margin', ':', f'{self.pricing.markup_percent}%']
            ]
            self._write_csv_header(writer, 'REKAP RAB', info_rows)
        """
        writer.writerow([title])
        writer.writerow([])  # Empty row
        
        for row in info_rows:
            writer.writerow(row)
        
        writer.writerow([])  # Empty row before table
    
    def _format_project_info(self) -> list:
        """
        Format project info sebagai list of rows untuk header.
        
        Returns:
            List of [label, ':', value]
        
        Usage:
            info_rows = self._format_project_info()
            self._write_csv_header(writer, 'REKAP RAB', info_rows)
        """
        return [
            ['Nama Project', ':', self.project.nama or '-'],
            ['Project ID', ':', str(self.project.id)],
            ['Kode Proyek', ':', self.project.index_project or '-'],
            ['Lokasi', ':', self.project.lokasi or '-'],
            ['Tahun Anggaran', ':', str(self.project.tahun_anggaran) if self.project.tahun_anggaran else '-'],
            ['Tanggal Export', ':', self.timestamp.strftime('%d-%m-%Y %H:%M')],
        ]
    
    # ========================================================================
    # PDF SHARED METHODS
    # ========================================================================
    
    def _create_pdf_doc(self, base_filename: str):
        """
        Create PDF document dengan landscape A4 dan proper margins.
        
        Args:
            base_filename: Base filename tanpa timestamp
        
        Returns:
            tuple: (buffer, doc) - BytesIO buffer dan SimpleDocTemplate instance
        
        Usage:
            buffer, doc = self._create_pdf_doc('rekap_rab.pdf')
            elements = []
            # ... add elements ...
            doc.build(elements)
        """
        if not REPORTLAB_AVAILABLE:
            raise ImportError("reportlab library not installed")
        
        buffer = BytesIO()
        
        # Add timestamp to title
        timestamp_str = self.timestamp.strftime('%Y%m%d_%H%M%S')
        title = f"{base_filename.replace('.pdf', '')}_{self.project.nama}_{timestamp_str}"
        
        doc = SimpleDocTemplate(
            buffer,
            pagesize=landscape(A4),
            rightMargin=ExportLayout.MARGIN_RIGHT * mm,
            leftMargin=ExportLayout.MARGIN_LEFT * mm,
            topMargin=ExportLayout.MARGIN_TOP * mm,
            bottomMargin=ExportLayout.MARGIN_BOTTOM * mm,
            title=title
        )
        
        return buffer, doc
    
    def _add_pdf_header(self, elements: list, title: str, subtitle: str = None):
        """
        Add standard PDF header dengan title, subtitle, dan project info.
        
        Args:
            elements: List untuk append elements
            title: Main title (e.g., 'RENCANA ANGGARAN BIAYA (RAB)')
            subtitle: Optional subtitle (default: project name)
        
        Usage:
            elements = []
            self._add_pdf_header(elements, 'REKAP RAB')
        """
        styles = getSampleStyleSheet()
        
        # Title
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=ExportFonts.TITLE,
            textColor=colors.HexColor(ExportColors.TEXT_PRIMARY),
            spaceAfter=3*mm,
            alignment=1,  # CENTER
            fontName='Helvetica-Bold',
            leading=22
        )
        elements.append(Paragraph(title, title_style))
        
        # Subtitle (project name by default)
        if subtitle is None:
            subtitle = self.project.nama or 'Nama Project'
        
        subtitle_style = ParagraphStyle(
            'CustomSubtitle',
            parent=styles['Heading2'],
            fontSize=ExportFonts.SUBTITLE,
            textColor=colors.HexColor(ExportColors.TEXT_PRIMARY),
            spaceAfter=10*mm,
            alignment=1,
            fontName='Helvetica-Bold',
            leading=20
        )
        elements.append(Paragraph(subtitle, subtitle_style))
        
        # Info table
        info_data = [
            ['Kode Proyek', ':', self.project.index_project or '-'],
            ['Lokasi', ':', self.project.lokasi or '-'],
            ['Tahun Anggaran', ':', str(self.project.tahun_anggaran) if self.project.tahun_anggaran else '-'],
        ]
        
        info_table = Table(info_data, colWidths=[40*mm, 5*mm, 80*mm])
        info_table.setStyle(TableStyle([
            ('FONTNAME', (0, 0), (-1, -1), 'Helvetica'),
            ('FONTSIZE', (0, 0), (-1, -1), ExportFonts.INFO),
            ('ALIGN', (0, 0), (0, -1), 'LEFT'),
            ('ALIGN', (1, 0), (1, -1), 'CENTER'),
            ('ALIGN', (2, 0), (2, -1), 'LEFT'),
            ('TEXTCOLOR', (0, 0), (-1, -1), colors.black),
            ('TOPPADDING', (0, 0), (-1, -1), 2),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 2),
        ]))
        
        elements.append(info_table)
        elements.append(Spacer(1, ExportLayout.SPACING_AFTER_HEADER*mm))
    
    def _create_pdf_table(self, table_data: list, col_widths: list) -> Table:
        """
        Create PDF table dengan standard styling.
        
        Args:
            table_data: List of lists (rows)
            col_widths: List of column widths in mm units
        
        Returns:
            Table instance dengan basic styling applied
        
        Usage:
            col_widths = ExportLayout.get_col_widths_mm(277)
            table = self._create_pdf_table(data, [
                col_widths['uraian'] * mm,
                col_widths['kode'] * mm,
                # ...
            ])
        """
        table = Table(table_data, colWidths=col_widths, repeatRows=1)
        
        # Basic table styling (header + borders)
        table_style = [
            # Header row styling
            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor(ExportColors.HEADER_BG)),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.HexColor(ExportColors.TEXT_PRIMARY)),
            ('FONTSIZE', (0, 0), (-1, 0), ExportFonts.HEADER),
            ('ALIGN', (0, 0), (-1, 0), 'CENTER'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 8),
            ('TOPPADDING', (0, 0), (-1, 0), 8),
            ('GRID', (0, 0), (-1, 0), 1, colors.black),
            
            # All cells
            ('FONTNAME', (0, 1), (-1, -1), 'Helvetica'),
            ('FONTSIZE', (0, 1), (-1, -1), ExportFonts.LEVEL3),
            ('GRID', (0, 0), (-1, -1), 0.5, colors.HexColor(ExportColors.BORDER_LIGHT)),
            ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
            ('TOPPADDING', (0, 1), (-1, -1), 6),
            ('BOTTOMPADDING', (0, 1), (-1, -1), 6),
            ('LEFTPADDING', (0, 1), (-1, -1), 6),
            ('RIGHTPADDING', (0, 1), (-1, -1), 6),
        ]
        
        table.setStyle(TableStyle(table_style))
        return table
    
    def _add_pdf_footer(self, elements: list, summary_text: str= None):
        """
        Add standard PDF footer dengan summary dan timestamp.
        
        Args:
            elements: List untuk append elements
            summary_text: Optional summary text (default: auto-generate)
        
        Usage:
            self._add_pdf_footer(elements, f'Total: {len(rows)} items')
        """
        elements.append(Spacer(1, ExportLayout.SPACING_AFTER_TABLE*mm))
        
        styles = getSampleStyleSheet()
        
        # Summary section
        if summary_text:
            summary_style = ParagraphStyle(
                'Summary',
                parent=styles['Normal'],
                fontSize=ExportFonts.SUMMARY_TEXT,
                textColor=colors.black,
                fontName='Helvetica-Bold'
            )
            elements.append(Paragraph(f"<b>RINGKASAN</b><br/>{summary_text}", summary_style))
            elements.append(Spacer(1, ExportLayout.SPACING_AFTER_SUMMARY*mm))
        
        # Footer timestamp
        footer_style = ParagraphStyle(
            'Footer',
            parent=styles['Normal'],
            fontSize=ExportFonts.TIMESTAMP,
            textColor=colors.HexColor(ExportColors.TEXT_SECONDARY),
            fontName='Helvetica-Oblique',
            alignment=2  # RIGHT
        )
        
        footer_text = f"Dicetak pada: {self.timestamp.strftime('%d-%m-%Y %H:%M')}"
        elements.append(Paragraph(footer_text, footer_style))
    
    # ========================================================================
    # WORD SHARED METHODS
    # ========================================================================
    
    def _create_word_doc(self):
        """
        Create Word document dengan landscape orientation.
        
        Returns:
            Document instance
        
        Usage:
            doc = self._create_word_doc()
            # ... add content ...
            buffer = BytesIO()
            doc.save(buffer)
        """
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx library not installed")
        
        doc = Document()
        
        # Set landscape orientation
        section = doc.sections[0]
        section.page_height = Inches(8.27)  # A4 width
        section.page_width = Inches(11.69)  # A4 height
        section.left_margin = Inches(ExportLayout.MARGIN_LEFT / 25.4)   # mm to inches
        section.right_margin = Inches(ExportLayout.MARGIN_RIGHT / 25.4)
        section.top_margin = Inches(ExportLayout.MARGIN_TOP / 25.4)
        section.bottom_margin = Inches(ExportLayout.MARGIN_BOTTOM / 25.4)
        
        return doc
    
    def _add_word_header(self, doc, title: str, subtitle: str = None):
        """
        Add standard Word header dengan title, subtitle, dan project info.
        
        Args:
            doc: Document instance
            title: Main title
            subtitle: Optional subtitle (default: project name)
        
        Usage:
            doc = self._create_word_doc()
            self._add_word_header(doc, 'REKAP RAB')
        """
        # Title
        title_para = doc.add_paragraph()
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title_para.add_run(title)
        title_run.font.bold = True
        title_run.font.size = Pt(ExportFonts.TITLE)
        title_run.font.name = 'Arial'
        
        # Subtitle (project name by default)
        if subtitle is None:
            subtitle = self.project.nama or 'Nama Project'
        
        subtitle_para = doc.add_paragraph()
        subtitle_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle_run = subtitle_para.add_run(subtitle)
        subtitle_run.font.bold = True
        subtitle_run.font.size = Pt(ExportFonts.SUBTITLE)
        subtitle_run.font.name = 'Arial'
        
        doc.add_paragraph()  # Spacing
        
        # Info table
        info_table = doc.add_table(rows=3, cols=3)
        info_table.alignment = WD_ALIGN_PARAGRAPH.LEFT
        
        info_data = [
            ['Kode Proyek', ':', self.project.index_project or '-'],
            ['Lokasi', ':', self.project.lokasi or '-'],
            ['Tahun Anggaran', ':', str(self.project.tahun_anggaran) if self.project.tahun_anggaran else '-'],
        ]
        
        for i, row_data in enumerate(info_data):
            row = info_table.rows[i]
            for j, cell_text in enumerate(row_data):
                cell = row.cells[j]
                cell.text = cell_text
                cell.paragraphs[0].runs[0].font.size = Pt(ExportFonts.INFO)
                cell.paragraphs[0].runs[0].font.name = 'Arial'
        
        # Set column widths
        info_table.columns[0].width = Inches(1.57)  # 40mm
        info_table.columns[1].width = Inches(0.20)  # 5mm
        info_table.columns[2].width = Inches(3.15)  # 80mm
        
        doc.add_paragraph()  # Spacing
    
    def _set_word_cell_color(self, cell, color_hex: str):
        """
        Set Word cell background color.
        
        Args:
            cell: Table cell
            color_hex: Color in hex format (e.g., 'e8e8e8')
        """
        shading = OxmlElement('w:shd')
        shading.set(qn('w:fill'), color_hex)
        cell._element.get_or_add_tcPr().append(shading)
    
    def _set_word_cell_border(self, cell, color='666666'):
        """
        Set Word cell borders.
        
        Args:
            cell: Table cell
            color: Border color in hex (default: '666666')
        """
        tc = cell._element
        tcPr = tc.get_or_add_tcPr()
        
        for edge in ('top', 'left', 'bottom', 'right'):
            edge_element = OxmlElement(f'w:{edge}')
            edge_element.set(qn('w:val'), 'single')
            edge_element.set(qn('w:sz'), '4')
            edge_element.set(qn('w:space'), '0')
            edge_element.set(qn('w:color'), color)
            tcPr.append(edge_element)
    
    # ========================================================================
    # ABSTRACT METHODS (to be implemented by subclasses)
    # ========================================================================
    
    def to_csv(self):
        """
        Export to CSV format.
        Must be implemented by subclass.
        
        Returns:
            HttpResponse with CSV file
        """
        raise NotImplementedError("Subclass must implement to_csv()")
    
    def to_pdf(self):
        """
        Export to PDF format.
        Must be implemented by subclass.
        
        Returns:
            HttpResponse with PDF file
        """
        raise NotImplementedError("Subclass must implement to_pdf()")
    
    def to_word(self):
        """
        Export to Word format.
        Must be implemented by subclass.
        
        Returns:
            HttpResponse with Word file
        """
        raise NotImplementedError("Subclass must implement to_word()")


# ============================================================================
# HELPER FUNCTIONS
# ============================================================================

def check_library_availability():
    """
    Check apakah reportlab dan python-docx tersedia.
    
    Returns:
        dict: {'reportlab': bool, 'docx': bool}
    
    Usage:
        libs = check_library_availability()
        if not libs['reportlab']:
            print("Install reportlab untuk PDF export")
    """
    return {
            'reportlab': REPORTLAB_AVAILABLE,
            'docx': DOCX_AVAILABLE,
        }


# ============================================================================
# CONFIG-BASED EXPORTER BASE (Phase 1, non-breaking)
# ============================================================================

try:
    # Import placed here to avoid circular imports during Django app loading
    from ..export_config import ExportConfig  # type: ignore
except Exception:  # pragma: no cover
    ExportConfig = None  # type: ignore


class ConfigExporterBase:
    """
    Base class for V2 exporters that operate with a single ExportConfig.
    Co-exists with legacy BaseExporter to avoid breaking existing exports.

    Provides:
    - self.config: immutable export configuration
    - _create_response(): unified HTTP response builder
    """

    def __init__(self, config):
        # duck-typed; prefer ExportConfig but avoid hard dependency at import time
        self.config = config

    def _create_response(self, content: bytes | bytearray | str, filename: str, content_type: str) -> HttpResponse:
        if isinstance(content, str):
            payload = content.encode('utf-8')
        else:
            payload = bytes(content)

        resp = HttpResponse(payload, content_type=content_type)
        resp['Content-Disposition'] = f'attachment; filename="{filename}"'
        return resp
