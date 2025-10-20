# ============================================================================
# FILE: detail_project/exports/rekap_rab.py
# ============================================================================
"""
Rekap RAB Exporter - Specific Implementation

Export Rekap RAB ke format CSV, PDF, dan Word dengan logic yang specific
untuk struktur Rekap RAB (Klasifikasi → SubKlasifikasi → Pekerjaan).

Architecture:
    RekapRABExporter (concrete)
    └── extends BaseExporter
        ├── to_csv() - implement CSV export
        ├── to_pdf() - implement PDF export
        └── to_word() - implement Word export

Usage:
    from .exports import RekapRABExporter
    
    exporter = RekapRABExporter(project, rekap_rows, pricing)
    
    # CSV
    response = exporter.to_csv()
    
    # PDF
    response = exporter.to_pdf()
    
    # Word
    response = exporter.to_word()

Author: Export Refactoring Phase 2
Created: 2025
"""

import csv
from io import BytesIO
from decimal import Decimal
from datetime import datetime

from django.http import HttpResponse

from .base import BaseExporter, REPORTLAB_AVAILABLE, DOCX_AVAILABLE
from ..export_config import (
    ExportColors,
    ExportFonts,
    ExportLayout,
    get_level_style,
    format_currency,
    format_volume,
)

# Conditional imports (sudah di-handle di base.py)
if REPORTLAB_AVAILABLE:
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import A4, landscape
    from reportlab.platypus import Table, TableStyle, Spacer
    from reportlab.lib.units import mm

if DOCX_AVAILABLE:
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH


# ============================================================================
# REKAP RAB EXPORTER
# ============================================================================

class RekapRABExporter(BaseExporter):
    """
    Exporter untuk Rekap RAB dengan support CSV, PDF, dan Word.
    
    Data format expected:
        List of dicts dengan keys:
        - level: 1 (Klasifikasi), 2 (Sub-Klasifikasi), 3 (Pekerjaan)
        - klasifikasi_nama: str
        - subklasifikasi_nama: str
        - pekerjaan_kode: str
        - pekerjaan_uraian: str
        - pekerjaan_satuan: str
        - volume: Decimal
        - A: Decimal (Tenaga Kerja)
        - B: Decimal (Bahan)
        - C: Decimal (Alat)
        - LAIN: Decimal (Lain-lain)
        - E_base: Decimal (Jumlah = A+B+C+LAIN)
        - F: Decimal (Profit/Margin)
        - G: Decimal (Harga Satuan = E_base + F)
        - total: Decimal (Total = G × volume)
    
    Usage:
        rekap_rows = compute_rekap_for_project(project)
        pricing = _get_or_create_pricing(project)
        
        exporter = RekapRABExporter(project, rekap_rows, pricing)
        response = exporter.to_csv()  # or to_pdf() or to_word()
    """
    
    # ========================================================================
    # CSV EXPORT
    # ========================================================================
    
    def to_csv(self):
        """
        Export Rekap RAB ke format CSV.
        
        Returns:
            HttpResponse dengan CSV file
        
        Format:
            - UTF-8 with BOM untuk Excel compatibility
            - Semicolon delimiter untuk Excel Indonesia
            - Header section dengan project info
            - Table dengan 15 kolom
            - Footer dengan total keseluruhan
            - Summary section
        """
        # Setup response
        response = self._create_csv_response('Rekap_RAB')
        writer = csv.writer(response, delimiter=';', quoting=csv.QUOTE_MINIMAL)
        
        # Markup percentage
        mp = self.pricing.markup_percent or Decimal("10.00")
        
        # ===== HEADER SECTION =====
        info_rows = self._format_project_info()
        info_rows.append(['Profit/Margin', ':', f'{mp}%'])
        
        self._write_csv_header(
            writer,
            'REKAP RENCANA ANGGARAN BIAYA (RAB)',
            info_rows
        )
        
        # ===== TABLE HEADER =====
        writer.writerow([
            'No',
            'Klasifikasi',
            'Sub-Klasifikasi',
            'Kode',
            'Uraian Pekerjaan',
            'Satuan',
            'Volume',
            'Tenaga Kerja (Rp)',
            'Bahan (Rp)',
            'Alat (Rp)',
            'Lain-lain (Rp)',
            'Jumlah (Rp)',
            'Profit/Margin (Rp)',
            'Harga Satuan (Rp)',
            'Total (Rp)'
        ])
        
        # ===== DATA ROWS =====
        total_all = Decimal("0")
        
        for idx, row in enumerate(self.data, start=1):
            writer.writerow([
                idx,
                row.get('klasifikasi_nama', '') or '-',
                row.get('subklasifikasi_nama', '') or '-',
                row.get('pekerjaan_kode', '') or '-',
                row.get('pekerjaan_uraian', '') or '-',
                row.get('pekerjaan_satuan', '') or '-',
                format_volume(row.get('volume')),
                format_currency(row.get('A')),
                format_currency(row.get('B')),
                format_currency(row.get('C')),
                format_currency(row.get('LAIN')),
                format_currency(row.get('E_base')),
                format_currency(row.get('F')),
                format_currency(row.get('G')),
                format_currency(row.get('total')),
            ])
            
            total_all += Decimal(str(row.get('total', 0)))
        
        # ===== FOOTER - TOTAL =====
        writer.writerow([])
        writer.writerow([
            '', '', '', '', '', '', '', '', '', '', '', '', '',
            'TOTAL KESELURUHAN:',
            format_currency(total_all)
        ])
        
        # ===== SUMMARY SECTION =====
        writer.writerow([])
        writer.writerow(['RINGKASAN'])
        writer.writerow(['Total Pekerjaan:', len(self.data)])
        writer.writerow(['Total Anggaran:', format_currency(total_all)])
        writer.writerow(['Profit/Margin:', f'{mp}%'])
        
        return response
    
    # ========================================================================
    # PDF EXPORT
    # ========================================================================
    
    def to_pdf(self):
        """
        Export Rekap RAB ke format PDF.
        
        Returns:
            HttpResponse dengan PDF file
        
        Format:
            - Landscape A4
            - Hierarchical styling (Level 1/2/3)
            - Color-coded rows
            - Professional layout
        """
        if not REPORTLAB_AVAILABLE:
            from django.http import JsonResponse
            return JsonResponse({
                'status': 'error',
                'message': 'Library reportlab belum terinstall.'
            }, status=500)
        
        # Setup PDF
        buffer, doc = self._create_pdf_doc('Rekap_RAB')
        elements = []
        
        # Add header
        self._add_pdf_header(
            elements,
            'RENCANA ANGGARAN BIAYA (RAB)',
            self.project.nama
        )
        
        # ===== PREPARE TABLE DATA =====
        table_data = [[
            'Uraian Pekerjaan',
            'Kode',
            'Satuan',
            'Volume',
            'Harga Satuan',
            'Total Harga'
        ]]
        
        total_all = Decimal('0')
        
        # Process each row dengan hierarki
        for row in self.data:
            level = row.get('level', 3)
            
            if level == 1:
                # KLASIFIKASI (Level 1)
                table_data.append([
                    row.get('klasifikasi_nama', ''),
                    '', '', '', '', ''
                ])
            elif level == 2:
                # SUB-KLASIFIKASI (Level 2)
                table_data.append([
                    '  ' + row.get('subklasifikasi_nama', ''),
                    '', '', '', '', ''
                ])
            else:
                # PEKERJAAN (Level 3)
                volume_str = format_volume(row.get('volume', 0))
                table_data.append([
                    '    ' + row.get('pekerjaan_uraian', ''),
                    row.get('pekerjaan_kode', '') or '-',
                    row.get('pekerjaan_satuan', '') or '-',
                    volume_str,
                    'Rp ' + format_currency(row.get('E_base', 0)),
                    'Rp ' + format_currency(row.get('total', 0))
                ])
                total_all += Decimal(str(row.get('total', 0)))
        
        # Footer - Total
        table_data.append([
            '', '', '', '',
            'TOTAL KESELURUHAN:',
            'Rp ' + format_currency(total_all)
        ])
        
        # ===== CREATE TABLE =====
        # Get column widths from config
        total_width = 277  # mm for A4 landscape
        col_widths_dict = ExportLayout.get_col_widths_mm(total_width)
        
        col_widths = [
            col_widths_dict['uraian'] * mm,
            col_widths_dict['kode'] * mm,
            col_widths_dict['satuan'] * mm,
            col_widths_dict['volume'] * mm,
            col_widths_dict['harga'] * mm,
            col_widths_dict['total'] * mm,
        ]
        
        table = self._create_pdf_table(table_data, col_widths)
        
        # ===== APPLY HIERARCHICAL STYLING =====
        # Get existing table style and add hierarchical styling
        additional_styles = []  # ✅ Create new list
        
        row_idx = 1  # Start after header
        for row in self.data:
            level = row.get('level', 3)
            style = get_level_style(level)
            
            # Apply background color
            additional_styles.extend([  # ✅ Use new list
                ('BACKGROUND', (0, row_idx), (-1, row_idx),
                 colors.HexColor(style['bg_color'])),
                ('FONTSIZE', (0, row_idx), (-1, row_idx), style['font_size']),
            ])
            
            # Apply bold font if needed
            if style['font_weight'] == 'bold':
                additional_styles.append(  # ✅ Use new list
                    ('FONTNAME', (0, row_idx), (-1, row_idx), 'Helvetica-Bold')
                )
            
            row_idx += 1
        
                # Footer row styling (TOTAL)
        additional_styles.extend([  # ✅ Use additional_styles
            ('BACKGROUND', (0, -1), (-1, -1), colors.HexColor(ExportColors.FOOTER_BG)),
            ('FONTNAME', (0, -1), (-1, -1), 'Helvetica-Bold'),
            ('FONTSIZE', (0, -1), (-1, -1), ExportFonts.FOOTER),
            ('LINEABOVE', (0, -1), (-1, -1), 2, colors.HexColor(ExportColors.BORDER_DARK)),
            ('TOPPADDING', (0, -1), (-1, -1), 8),
            ('BOTTOMPADDING', (0, -1), (-1, -1), 8),
            ('ALIGN', (4, -1), (5, -1), 'RIGHT'),
        ])
        
        # Apply numeric column alignment (for pekerjaan rows)
        additional_styles.extend([  # ✅ Use additional_styles
            ('ALIGN', (3, 1), (3, -2), 'RIGHT'),  # Volume (exclude header and footer)
            ('ALIGN', (4, 1), (4, -2), 'RIGHT'),  # Harga
            ('ALIGN', (5, 1), (5, -2), 'RIGHT'),  # Total
        ])
        
        # ✅ Apply additional styles to table
        table.setStyle(TableStyle(additional_styles))
        elements.append(table)

        # ===== ADD FOOTER =====
        mp = self.pricing.markup_percent or Decimal("10.00")
        summary_text = f"Total Pekerjaan: {len(self.data)} | Total Anggaran: Rp {format_currency(total_all)} | Profit/Margin: {mp}%"
        self._add_pdf_footer(elements, summary_text)
        
        # Build PDF
        doc.build(elements)
        
        buffer.seek(0)
        response = HttpResponse(buffer.getvalue(), content_type='application/pdf')
        
        timestamp_str = self.timestamp.strftime('%Y%m%d_%H%M%S')
        filename = f'Rekap_RAB_{self.project.nama}_{timestamp_str}.pdf'
        response['Content-Disposition'] = f'attachment; filename="{filename}"'
        
        return response
    
    # ========================================================================
    # WORD EXPORT
    # ========================================================================
    
    def to_word(self):
        """
        Export Rekap RAB ke format Word.
        
        Returns:
            HttpResponse dengan Word file
        
        Format:
            - Landscape orientation
            - Hierarchical styling
            - Professional layout
        """
        if not DOCX_AVAILABLE:
            from django.http import JsonResponse
            return JsonResponse({
                'status': 'error',
                'message': 'Library python-docx belum terinstall.'
            }, status=500)
        
        # Setup document
        doc = self._create_word_doc()
        
        # Add header
        self._add_word_header(
            doc,
            'RENCANA ANGGARAN BIAYA (RAB)',
            self.project.nama
        )
        
        # ===== CREATE TABLE =====
        num_rows = 1 + len(self.data) + 1  # header + data + footer
        table = doc.add_table(rows=num_rows, cols=6)
        table.style = 'Table Grid'
        
        # Calculate content width
        section = doc.sections[0]
        page_width = section.page_width.inches
        margin_left = section.left_margin.inches
        margin_right = section.right_margin.inches
        content_width = page_width - margin_left - margin_right
        
        # Get widths from config
        col_widths = ExportLayout.get_col_widths_inches(content_width)
        
        from docx.shared import Inches
        table.columns[0].width = Inches(col_widths['uraian'])
        table.columns[1].width = Inches(col_widths['kode'])
        table.columns[2].width = Inches(col_widths['satuan'])
        table.columns[3].width = Inches(col_widths['volume'])
        table.columns[4].width = Inches(col_widths['harga'])
        table.columns[5].width = Inches(col_widths['total'])
        
        # ===== HEADER ROW =====
        header_cells = table.rows[0].cells
        headers = ['Uraian Pekerjaan', 'Kode', 'Satuan', 'Volume', 'Harga Satuan', 'Total Harga']
        
        for i, header_text in enumerate(headers):
            cell = header_cells[i]
            cell.text = header_text
            
            # Styling
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.font.bold = True
                    run.font.size = Pt(ExportFonts.HEADER)
                    run.font.name = 'Arial'
            
            self._set_word_cell_color(cell, ExportColors.HEADER_BG)
            self._set_word_cell_border(cell, color='000000')
        
        # ===== DATA ROWS =====
        total_all = Decimal('0')
        
        for idx, row_data in enumerate(self.data):
            row = table.rows[idx + 1]
            cells = row.cells
            level = row_data.get('level', 3)
            style = get_level_style(level)
            
            # Set text with indentation
            if level == 1:
                # KLASIFIKASI
                cells[0].text = row_data.get('klasifikasi_nama', '')
                # Clear other cells
                for i in range(1, 6):
                    cells[i].text = ''
            elif level == 2:
                # SUB-KLASIFIKASI
                cells[0].text = style['indent_text'] + row_data.get('subklasifikasi_nama', '')
                # Clear other cells
                for i in range(1, 6):
                    cells[i].text = ''
            else:  # level 3
                # PEKERJAAN - populate all columns
                cells[0].text = style['indent_text'] + row_data.get('pekerjaan_uraian', '')
                cells[1].text = row_data.get('pekerjaan_kode', '') or '-'
                cells[2].text = row_data.get('pekerjaan_satuan', '') or '-'
                cells[3].text = format_volume(row_data.get('volume', 0))
                cells[4].text = 'Rp ' + format_currency(row_data.get('E_base', 0))
                cells[5].text = 'Rp ' + format_currency(row_data.get('total', 0))
                total_all += Decimal(str(row_data.get('total', 0)))
            
            # Apply styling to all cells
            for cell in cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.font.size = Pt(style['font_size'])
                        run.font.name = 'Arial'
                        if style['font_weight'] == 'bold':
                            run.font.bold = True
                self._set_word_cell_color(cell, style['bg_color'])
                self._set_word_cell_border(cell, color='666666')
        
        # ===== FOOTER ROW (TOTAL) =====
        footer_row = table.rows[-1]
        cells = footer_row.cells
        
        cells[0].text = ''
        cells[1].text = ''
        cells[2].text = ''
        cells[3].text = ''
        cells[4].text = 'TOTAL KESELURUHAN:'
        cells[5].text = 'Rp ' + format_currency(total_all)
        
        # Style footer
        for i, cell in enumerate(cells):
            for paragraph in cell.paragraphs:
                if i >= 4:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                for run in paragraph.runs:
                    run.font.bold = True
                    run.font.size = Pt(ExportFonts.FOOTER)
                    run.font.name = 'Arial'
            self._set_word_cell_color(cell, ExportColors.FOOTER_BG)
            self._set_word_cell_border(cell, color='000000')
        
        # ===== SUMMARY SECTION =====
        doc.add_paragraph()
        
        summary = doc.add_paragraph()
        summary_run = summary.add_run('RINGKASAN')
        summary_run.font.bold = True
        summary_run.font.size = Pt(ExportFonts.SUMMARY)
        summary_run.font.name = 'Arial'
        
        mp = self.pricing.markup_percent or Decimal("10.00")
        summary_details = doc.add_paragraph()
        summary_text = f"Total Pekerjaan: {len(self.data)} | Total Anggaran: Rp {format_currency(total_all)} | Profit/Margin: {mp}%"
        summary_run = summary_details.add_run(summary_text)
        summary_run.font.size = Pt(ExportFonts.SUMMARY_TEXT)
        summary_run.font.name = 'Arial'
        
        # ===== FOOTER =====
        doc.add_paragraph()
        footer = doc.add_paragraph()
        footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        footer_run = footer.add_run(f"Dicetak pada: {self.timestamp.strftime('%d-%m-%Y %H:%M')}")
        footer_run.font.size = Pt(ExportFonts.TIMESTAMP)
        footer_run.font.italic = True
        
        from docx.shared import RGBColor
        footer_run.font.color.rgb = RGBColor(102, 102, 102)
        footer_run.font.name = 'Arial'
        
        # ===== SAVE TO BUFFER =====
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        
        # ===== RESPONSE =====
        response = HttpResponse(
            buffer.getvalue(),
            content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        
        timestamp_str = self.timestamp.strftime('%Y%m%d_%H%M%S')
        filename = f'Rekap_RAB_{self.project.nama}_{timestamp_str}.docx'
        response['Content-Disposition'] = f'attachment; filename="{filename}"'
        
        return response