# ============================================================================
# FILE: detail_project/exports/rekap_kebutuhan.py
# ============================================================================
"""
Rekap Kebutuhan Exporter - Specific Implementation

Export Rekap Kebutuhan (agregasi material/tenaga/alat) ke format CSV, PDF, dan Word.

Data Structure:
    Flat table (no hierarchy) dengan agregasi quantity per item.
    Columns: Kategori, Kode, Uraian, Satuan, Quantity

Architecture:
    RekapKebutuhanExporter (concrete)
    └── extends BaseExporter
        ├── to_csv() - implement CSV export
        ├── to_pdf() - implement PDF export (NEW!)
        └── to_word() - implement Word export (NEW!)

Usage:
    from .exports import RekapKebutuhanExporter
    from .services import compute_kebutuhan_items
    
    rows = compute_kebutuhan_items(project)
    exporter = RekapKebutuhanExporter(project, rows, pricing)
    
    # CSV
    response = exporter.to_csv()
    
    # PDF (NEW!)
    response = exporter.to_pdf()
    
    # Word (NEW!)
    response = exporter.to_word()

Author: Export Refactoring Phase 3
Created: 2025
"""

import csv
from io import BytesIO
from decimal import Decimal

from django.http import HttpResponse

from .base import BaseExporter, REPORTLAB_AVAILABLE, DOCX_AVAILABLE
from ..export_config import (
    ExportColors,
    ExportFonts,
    ExportLayout,
    format_currency,
    format_volume,
)

# Conditional imports
if REPORTLAB_AVAILABLE:
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import A4, landscape
    from reportlab.platypus import Table, TableStyle
    from reportlab.lib.units import mm

if DOCX_AVAILABLE:
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH


# ============================================================================
# REKAP KEBUTUHAN EXPORTER
# ============================================================================

class RekapKebutuhanExporter(BaseExporter):
    """
    Exporter untuk Rekap Kebutuhan dengan support CSV, PDF, dan Word.
    
    Data format expected:
        List of dicts dengan keys:
        - kategori: str (TK/BHN/ALT/LAIN)
        - kode: str
        - uraian: str
        - satuan: str
        - quantity: Decimal (agregasi dari semua pekerjaan)
    
    Note: 
        Rekap Kebutuhan adalah flat table (no hierarchy), lebih simple
        dibanding Rekap RAB.
    
    Usage:
        from .services import compute_kebutuhan_items
        
        rows = compute_kebutuhan_items(project)
        exporter = RekapKebutuhanExporter(project, rows, None)  # no pricing needed
        response = exporter.to_csv()  # or to_pdf() or to_word()
    """
    
    # ========================================================================
    # CSV EXPORT
    # ========================================================================
    
    def to_csv(self):
        """
        Export Rekap Kebutuhan ke format CSV.
        
        Returns:
            HttpResponse dengan CSV file
        
        Format:
            - UTF-8 with BOM untuk Excel compatibility
            - Semicolon delimiter untuk Excel Indonesia
            - Header section dengan project info
            - Simple table: Kategori, Kode, Uraian, Satuan, Quantity
            - Summary section dengan counts per kategori
        """
        # Setup response
        response = self._create_csv_response('Rekap_Kebutuhan')
        writer = csv.writer(response, delimiter=';', quoting=csv.QUOTE_MINIMAL)
        
        # ===== HEADER SECTION =====
        info_rows = self._format_project_info()
        
        self._write_csv_header(
            writer,
            'REKAP KEBUTUHAN MATERIAL, TENAGA KERJA, DAN ALAT',
            info_rows
        )
        
        # ===== TABLE HEADER =====
        writer.writerow([
            'Kategori',
            'Kode',
            'Uraian',
            'Satuan',
            'Quantity'
        ])
        
        # ===== DATA ROWS =====
        # Count per kategori
        counts = {"TK": 0, "BHN": 0, "ALT": 0, "LAIN": 0}
        
        for row in self.data:
            kategori = row.get('kategori', '')
            
            writer.writerow([
                kategori,
                row.get('kode', '') or '-',
                row.get('uraian', '') or '-',
                row.get('satuan', '') or '-',
                format_volume(row.get('quantity', 0))
            ])
            
            # Count
            if kategori in counts:
                counts[kategori] += 1
        
        # ===== SUMMARY SECTION =====
        writer.writerow([])
        writer.writerow(['RINGKASAN'])
        writer.writerow(['Total Items:', len(self.data)])
        writer.writerow(['Tenaga Kerja:', counts['TK']])
        writer.writerow(['Bahan:', counts['BHN']])
        writer.writerow(['Alat:', counts['ALT']])
        writer.writerow(['Lain-lain:', counts['LAIN']])
        
        return response
    
    # ========================================================================
    # PDF EXPORT (NEW!)
    # ========================================================================
    
    def to_pdf(self):
        """
        Export Rekap Kebutuhan ke format PDF.
        
        Returns:
            HttpResponse dengan PDF file
        
        Format:
            - Landscape A4
            - Simple table (5 columns, no hierarchy)
            - Summary dengan counts per kategori
        """
        if not REPORTLAB_AVAILABLE:
            from django.http import JsonResponse
            return JsonResponse({
                'status': 'error',
                'message': 'Library reportlab belum terinstall.'
            }, status=500)
        
        # Setup PDF
        buffer, doc = self._create_pdf_doc('Rekap_Kebutuhan')
        elements = []
        
        # Add header
        self._add_pdf_header(
            elements,
            'REKAP KEBUTUHAN MATERIAL, TENAGA KERJA, DAN ALAT',
            self.project.nama
        )
        
        # ===== PREPARE TABLE DATA =====
        table_data = [[
            'Kategori',
            'Kode',
            'Uraian',
            'Satuan',
            'Quantity'
        ]]
        
        # Count per kategori
        counts = {"TK": 0, "BHN": 0, "ALT": 0, "LAIN": 0}
        
        # Process each row
        for row in self.data:
            kategori = row.get('kategori', '')
            
            table_data.append([
                kategori,
                row.get('kode', '') or '-',
                row.get('uraian', '') or '-',
                row.get('satuan', '') or '-',
                format_volume(row.get('quantity', 0))
            ])
            
            # Count
            if kategori in counts:
                counts[kategori] += 1
        
        # ===== CREATE TABLE =====
        # Column widths (simpler than Rekap RAB - only 5 columns)
        total_width = 277  # mm for A4 landscape
        col_widths = [
            0.12 * total_width * mm,  # Kategori
            0.15 * total_width * mm,  # Kode
            0.45 * total_width * mm,  # Uraian (largest)
            0.13 * total_width * mm,  # Satuan
            0.15 * total_width * mm,  # Quantity
        ]
        
        table = self._create_pdf_table(table_data, col_widths)
        
        # Additional styling for kebutuhan table
        table_style_additions = [
            # Align quantity to right
            ('ALIGN', (4, 1), (4, -1), 'RIGHT'),
            
            # Alternate row colors for better readability
            ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.HexColor('f9f9f9')]),
        ]
        
        existing_style = table._cellStyles
        existing_style.extend(table_style_additions)
        table.setStyle(TableStyle(existing_style))
        
        elements.append(table)
        
        # ===== ADD FOOTER WITH SUMMARY =====
        summary_text = (
            f"Total Items: {len(self.data)} | "
            f"Tenaga Kerja: {counts['TK']} | "
            f"Bahan: {counts['BHN']} | "
            f"Alat: {counts['ALT']} | "
            f"Lain-lain: {counts['LAIN']}"
        )
        self._add_pdf_footer(elements, summary_text)
        
        # Build PDF
        doc.build(elements)
        
        buffer.seek(0)
        response = HttpResponse(buffer.getvalue(), content_type='application/pdf')
        
        timestamp_str = self.timestamp.strftime('%Y%m%d_%H%M%S')
        filename = f'Rekap_Kebutuhan_{self.project.nama}_{timestamp_str}.pdf'
        response['Content-Disposition'] = f'attachment; filename="{filename}"'
        
        return response
    
    # ========================================================================
    # WORD EXPORT (NEW!)
    # ========================================================================
    
    def to_word(self):
        """
        Export Rekap Kebutuhan ke format Word.
        
        Returns:
            HttpResponse dengan Word file
        
        Format:
            - Landscape orientation
            - Simple table (5 columns)
            - Summary section
        """
        if not DOCX_AVAILABLE:
            from django.http import JsonResponse
            return JsonResponse({
                'status': 'error',
                'message': 'Library python-docx belum terinstall.'
            }, status=500)
        
        # Setup document
        doc = self._create_word_doc()
        
        # Add header
        self._add_word_header(
            doc,
            'REKAP KEBUTUHAN MATERIAL, TENAGA KERJA, DAN ALAT',
            self.project.nama
        )
        
        # ===== CREATE TABLE =====
        num_rows = 1 + len(self.data)  # header + data (no footer for kebutuhan)
        table = doc.add_table(rows=num_rows, cols=5)
        table.style = 'Table Grid'
        
        # Calculate content width
        section = doc.sections[0]
        page_width = section.page_width.inches
        margin_left = section.left_margin.inches
        margin_right = section.right_margin.inches
        content_width = page_width - margin_left - margin_right
        
        # Set column widths (same proportions as PDF)
        table.columns[0].width = Inches(0.12 * content_width)  # Kategori
        table.columns[1].width = Inches(0.15 * content_width)  # Kode
        table.columns[2].width = Inches(0.45 * content_width)  # Uraian
        table.columns[3].width = Inches(0.13 * content_width)  # Satuan
        table.columns[4].width = Inches(0.15 * content_width)  # Quantity
        
        # ===== HEADER ROW =====
        header_cells = table.rows[0].cells
        headers = ['Kategori', 'Kode', 'Uraian', 'Satuan', 'Quantity']
        
        for i, header_text in enumerate(headers):
            cell = header_cells[i]
            cell.text = header_text
            
            # Styling
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.font.bold = True
                    run.font.size = Pt(ExportFonts.HEADER)
                    run.font.name = 'Arial'
            
            self._set_word_cell_color(cell, ExportColors.HEADER_BG)
            self._set_word_cell_border(cell, color='000000')
        
        # ===== DATA ROWS =====
        counts = {"TK": 0, "BHN": 0, "ALT": 0, "LAIN": 0}
        
        for idx, row_data in enumerate(self.data):
            row = table.rows[idx + 1]
            cells = row.cells
            
            kategori = row_data.get('kategori', '')
            
            # Populate cells
            cells[0].text = kategori
            cells[1].text = row_data.get('kode', '') or '-'
            cells[2].text = row_data.get('uraian', '') or '-'
            cells[3].text = row_data.get('satuan', '') or '-'
            cells[4].text = format_volume(row_data.get('quantity', 0))
            
            # Count
            if kategori in counts:
                counts[kategori] += 1
            
            # Apply styling to all cells
            for i, cell in enumerate(cells):
                for para in cell.paragraphs:
                    # Align quantity to right
                    if i == 4:
                        para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    
                    for run in para.runs:
                        run.font.size = Pt(ExportFonts.LEVEL3)
                        run.font.name = 'Arial'
                
                # Alternate row colors
                if idx % 2 == 0:
                    self._set_word_cell_color(cell, 'ffffff')
                else:
                    self._set_word_cell_color(cell, 'f9f9f9')
                
                self._set_word_cell_border(cell, color='666666')
        
        # ===== SUMMARY SECTION =====
        doc.add_paragraph()
        
        summary = doc.add_paragraph()
        summary_run = summary.add_run('RINGKASAN')
        summary_run.font.bold = True
        summary_run.font.size = Pt(ExportFonts.SUMMARY)
        summary_run.font.name = 'Arial'
        
        summary_details = doc.add_paragraph()
        summary_text = (
            f"Total Items: {len(self.data)} | "
            f"Tenaga Kerja: {counts['TK']} | "
            f"Bahan: {counts['BHN']} | "
            f"Alat: {counts['ALT']} | "
            f"Lain-lain: {counts['LAIN']}"
        )
        summary_run = summary_details.add_run(summary_text)
        summary_run.font.size = Pt(ExportFonts.SUMMARY_TEXT)
        summary_run.font.name = 'Arial'
        
        # ===== FOOTER =====
        doc.add_paragraph()
        footer = doc.add_paragraph()
        footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        footer_run = footer.add_run(f"Dicetak pada: {self.timestamp.strftime('%d-%m-%Y %H:%M')}")
        footer_run.font.size = Pt(ExportFonts.TIMESTAMP)
        footer_run.font.italic = True
        
        from docx.shared import RGBColor
        footer_run.font.color.rgb = RGBColor(102, 102, 102)
        footer_run.font.name = 'Arial'
        
        # ===== SAVE TO BUFFER =====
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        
        # ===== RESPONSE =====
        response = HttpResponse(
            buffer.getvalue(),
            content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        
        timestamp_str = self.timestamp.strftime('%Y%m%d_%H%M%S')
        filename = f'Rekap_Kebutuhan_{self.project.nama}_{timestamp_str}.docx'
        response['Content-Disposition'] = f'attachment; filename="{filename}"'
        
        return response