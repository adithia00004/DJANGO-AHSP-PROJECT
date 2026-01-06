# ============================================================================
# FILE: detail_project/exports/word_exporter.py
# ============================================================================
"""
Word Document Export Handler

Exports Jadwal Pekerjaan data to Microsoft Word (.docx) format.
Supports 3 report types:
- Rekap: Full summary with Grid Planned/Actual + Kurva S
- Bulanan (Monthly): Progress report with Kurva S
- Mingguan (Weekly): Compact weekly progress

Uses python-docx for native Word element manipulation.

Author: Word Export Phase 1
Created: 2025
"""

from io import BytesIO
from typing import Dict, Any, List
from django.http import HttpResponse

from docx import Document
from docx.shared import Inches, Mm, Pt, Cm, RGBColor
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from ..export_config import (
    ExportConfig,
    ExportColors,
    ExportFonts,
    JadwalExportLayout,
    get_level_style,
    build_identity_rows,
)
from .table_styles import UnifiedTableStyles as UTS


class WordExporter:
    """
    Word Document Exporter for Jadwal Pekerjaan reports.
    
    Usage:
        config = ExportConfig(...)
        exporter = WordExporter(config)
        response = exporter.export_rekap(data)
    """
    
    def __init__(self, config: ExportConfig):
        """
        Initialize Word exporter with configuration.
        
        Args:
            config: ExportConfig with project info and styling settings
        """
        self.config = config
        self.doc = None
    
    # =========================================================================
    # PUBLIC EXPORT METHODS
    # =========================================================================
    
    def export_professional(self, data: Dict[str, Any]) -> HttpResponse:
        """
        Export professional formatted Word document.
        
        This is the main entry point called by ExportManager.
        Dispatches to appropriate export method based on report_type.
        
        Args:
            data: Export data containing report_type and content
            
        Returns:
            HttpResponse with .docx file
        """
        report_type = data.get('report_type', 'rekap')
        
        if report_type == 'monthly':
            return self.export_monthly(data)
        elif report_type == 'weekly':
            return self.export_weekly(data)
        else:
            return self.export_rekap(data)
    
    def export(self, data: Dict[str, Any]) -> HttpResponse:
        """
        Generic export method for backward compatibility.
        
        Called by ExportManager for non-professional exports like rekap_rab.
        Creates a simple document with pages of tables.
        
        Args:
            data: Export data with 'pages' list
            
        Returns:
            HttpResponse with .docx file
        """
        # Check if this is Rincian AHSP data
        sections = data.get('sections', [])
        is_rincian_ahsp = bool(
            sections and 
            isinstance(sections[0], dict) and 
            'pekerjaan' in sections[0] and 
            'groups' in sections[0]
        )

        if is_rincian_ahsp:
            return self._export_rincian_ahsp(data)

        self.doc = Document()
        self._setup_page_layout('A4', 'portrait')
        
        # Handle single-table data (e.g., Harga Items) vs multi-page data
        pages = data.get('pages', [])
        if not pages and 'table_data' in data:
            # Single table data - wrap it as a single page
            pages = [data]
        
        for idx, page in enumerate(pages):
            # Page title
            title = page.get('title') or self.config.title or f'Page {idx + 1}'
            self._build_section_header(title)
            
            # Project identity
            from ..export_config import build_identity_rows
            for label, _, value in build_identity_rows(self.config):
                para = self.doc.add_paragraph()
                para.add_run(f"{label}: ").bold = True
                para.add_run(str(value))
            
            self.doc.add_paragraph()  # Spacing
            
            # Build table from page data
            table_data = page.get('table_data', {})
            headers = table_data.get('headers', [])
            rows = table_data.get('rows', [])
            row_types = page.get('row_types', [])
            is_pengesahan_page = page.get('include_signatures', False)
            
            if headers and rows:
                if is_pengesahan_page:
                    # Pengesahan page - use dedicated 3-column table
                    self._build_pengesahan_word_table(table_data, page.get('col_widths', []))
                else:
                    # Regular table - use existing logic
                    num_cols = len(headers)
                    num_rows = len(rows) + 1
                    
                    table = self.doc.add_table(rows=num_rows, cols=num_cols)
                    table.style = 'Table Grid'
                
                    # Header row
                    for col_idx, header_text in enumerate(headers):
                        cell = table.rows[0].cells[col_idx]
                        cell.text = str(header_text)
                        self._style_header_cell(cell)
                    
                    # Data rows
                    for row_idx, row_data in enumerate(rows):
                        row_type = row_types[row_idx] if row_idx < len(row_types) else 'item'
                        table_row = table.rows[row_idx + 1]
                        
                        if row_type == 'category':
                            # Category row - merge all cells and bold
                            # Merge cells for category header
                            for col_idx in range(1, num_cols):
                                table_row.cells[0].merge(table_row.cells[col_idx])
                            table_row.cells[0].text = str(row_data[0]) if row_data else ''
                            for para in table_row.cells[0].paragraphs:
                                for run in para.runs:
                                    run.bold = True
                                para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                            # Apply gray background using shading
                            from docx.oxml.ns import qn
                            from docx.oxml import OxmlElement
                            shading = OxmlElement('w:shd')
                            shading.set(qn('w:fill'), 'E8E8E8')
                            table_row.cells[0]._tc.get_or_add_tcPr().append(shading)
                        else:
                            # Normal item row
                            for col_idx, cell_value in enumerate(row_data):
                                if col_idx < num_cols:
                                    cell = table_row.cells[col_idx]
                                    text = str(cell_value) if cell_value else ''
                                    # Handle multi-line text (with \n)
                                    if '\n' in text:
                                        lines = text.split('\n')
                                        cell.text = ''  # Clear default paragraph
                                        for i, line in enumerate(lines):
                                            if i == 0:
                                                cell.paragraphs[0].text = line
                                            else:
                                                cell.add_paragraph(line)
                                    else:
                                        cell.text = text
                    
                    self._enable_header_repeat(table)
                    
                    # Set column widths if specified
                    col_widths = page.get('col_widths', [])
                    if col_widths:
                        for row in table.rows:
                            for col_idx, cell in enumerate(row.cells):
                                if col_idx < len(col_widths):
                                    cell.width = Mm(col_widths[col_idx])
            
            # Footer rows
            footer_rows = page.get('footer_rows', [])
            if footer_rows:
                self.doc.add_paragraph()  # Spacing
                for footer in footer_rows:
                    para = self.doc.add_paragraph()
                    if isinstance(footer, (list, tuple)) and len(footer) >= 2:
                        para.add_run(f"{footer[0]}: ").bold = True
                        para.add_run(str(footer[1]))
                    else:
                        para.add_run(str(footer))
            
            # Add signatures if this page has include_signatures=True
            if page.get('include_signatures') and self.config.signature_config.enabled:
                self.doc.add_paragraph()  # Spacing
                self._build_signature_section()
            
            # Page break if not last page
            if idx < len(pages) - 1:
                self.doc.add_page_break()
        
        return self._create_response('export')

    def _export_rincian_ahsp(self, data: Dict[str, Any]) -> HttpResponse:
        """
        Export Rincian AHSP to Word document.
        
        Structure:
        1. Rekap (summary table)
        2. Rincian (detail per pekerjaan)
        3. Lembar Pengesahan (at bottom of last rincian page)
        """
        self.doc = Document()
        self._setup_page_layout('A4', 'portrait')
        
        sections = data.get('sections', [])
        
        # ========== SECTION 1: REKAP ==========
        title_para = self.doc.add_paragraph()
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title_para.add_run('REKAP ANALISA HARGA SATUAN PEKERJAAN')
        title_run.bold = True
        title_run.font.size = Pt(16)
        
        # Identity rows
        identity_rows = build_identity_rows(self.config)
        if identity_rows:
            self.doc.add_paragraph()
            id_table = self.doc.add_table(rows=len(identity_rows), cols=3)
            for i, row_data in enumerate(identity_rows):
                for j, cell_text in enumerate(row_data):
                    id_table.rows[i].cells[j].text = str(cell_text)
        
        self.doc.add_paragraph()
        
        # Rekap table
        rekap_headers = ['No', 'Kode', 'Uraian Pekerjaan', 'E — Jumlah', 'F — Profit/Margin', 'G — Harga Satuan']
        rekap_table = self.doc.add_table(rows=len(sections) + 1, cols=6)
        rekap_table.style = 'Table Grid'
        
        # Header row
        for col_idx, header_text in enumerate(rekap_headers):
            cell = rekap_table.rows[0].cells[col_idx]
            cell.text = header_text
            self._style_header_cell(cell)
        
        # Data rows
        for idx, section in enumerate(sections):
            pekerjaan = section.get('pekerjaan', {})
            totals = section.get('totals', {})
            row = rekap_table.rows[idx + 1]
            
            row.cells[0].text = str(idx + 1)
            row.cells[1].text = pekerjaan.get('kode', '')
            row.cells[2].text = pekerjaan.get('uraian', '')
            row.cells[3].text = totals.get('E', '0')
            row.cells[4].text = totals.get('F', '0')
            row.cells[5].text = totals.get('G', '0')
            # Bold G column
            for para in row.cells[5].paragraphs:
                for run in para.runs:
                    run.bold = True
        
        # Column widths for rekap
        rekap_widths = [Mm(8), Mm(20), Mm(55), Mm(30), Mm(30), Mm(32)]
        for row in rekap_table.rows:
            for col_idx, cell in enumerate(row.cells):
                cell.width = rekap_widths[col_idx]
        
        self.doc.add_page_break()
        
        # ========== SECTION 2: RINCIAN ==========
        rincian_title = self.doc.add_paragraph()
        rincian_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        rincian_run = rincian_title.add_run('RINCIAN ANALISA HARGA SATUAN PEKERJAAN')
        rincian_run.bold = True
        rincian_run.font.size = Pt(16)
        
        self.doc.add_paragraph()
        
        # Write each pekerjaan section (no Grand Total)
        for idx, section in enumerate(sections):
            pekerjaan = section.get('pekerjaan', {})
            groups = section.get('groups', [])
            totals = section.get('totals', {})
            
            pek_kode = pekerjaan.get('kode', '')
            pek_uraian = pekerjaan.get('uraian', '')
            
            # Pekerjaan header
            header_para = self.doc.add_paragraph()
            header_run = header_para.add_run(f"{pek_kode} - {pek_uraian}" if pek_kode else pek_uraian)
            header_run.bold = True
            header_run.font.size = Pt(11)
            
            # Detail table headers
            headers = ['No', 'Uraian', 'Kode', 'Satuan', 'Koefisien', 'Harga Satuan', 'Jumlah Harga']
            
            # Count total rows for table
            total_rows = 1  # header
            for group in groups:
                if group.get('rows'):
                    total_rows += 1  # group title
                    total_rows += len(group.get('rows', []))
                    total_rows += 1  # subtotal
            total_rows += 3  # E, F, G
            
            table = self.doc.add_table(rows=total_rows, cols=7)
            table.style = 'Table Grid'
            
            # Header row
            for col_idx, header_text in enumerate(headers):
                cell = table.rows[0].cells[col_idx]
                cell.text = header_text
                self._style_header_cell(cell)
            
            row_idx = 1
            
            # Write groups
            for group in groups:
                group_title = group.get('title', '')
                group_rows = group.get('rows', [])
                group_subtotal = group.get('subtotal', '')
                
                if not group_rows:
                    continue
                
                # Group title
                if row_idx < len(table.rows):
                    table.rows[row_idx].cells[0].text = group_title
                    for para in table.rows[row_idx].cells[0].paragraphs:
                        for run in para.runs:
                            run.bold = True
                            run.italic = True
                            run.font.size = Pt(9)
                    row_idx += 1
                
                # Group detail rows
                for row_data in group_rows:
                    if row_idx < len(table.rows):
                        for col_idx, val in enumerate(row_data):
                            if col_idx < 7:
                                table.rows[row_idx].cells[col_idx].text = str(val) if val else ''
                        row_idx += 1
                
                # Subtotal row
                if row_idx < len(table.rows):
                    table.rows[row_idx].cells[0].text = f"Subtotal {group.get('short_title', '')}"
                    for para in table.rows[row_idx].cells[0].paragraphs:
                        para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                        for run in para.runs:
                            run.bold = True
                    table.rows[row_idx].cells[6].text = str(group_subtotal)
                    for para in table.rows[row_idx].cells[6].paragraphs:
                        for run in para.runs:
                            run.bold = True
                    row_idx += 1
            
            # Total E
            if row_idx < len(table.rows):
                table.rows[row_idx].cells[0].text = "Jumlah (E)"
                for para in table.rows[row_idx].cells[0].paragraphs:
                    para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    for run in para.runs:
                        run.bold = True
                table.rows[row_idx].cells[6].text = totals.get('E', '0')
                for para in table.rows[row_idx].cells[6].paragraphs:
                    for run in para.runs:
                        run.bold = True
                row_idx += 1
            
            # Total F (Profit/Margin)
            if row_idx < len(table.rows):
                markup = totals.get('markup_eff', '10.00')
                table.rows[row_idx].cells[0].text = f"Profit/Margin {markup}% (F)"
                for para in table.rows[row_idx].cells[0].paragraphs:
                    para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    for run in para.runs:
                        run.bold = True
                table.rows[row_idx].cells[6].text = totals.get('F', '0')
                for para in table.rows[row_idx].cells[6].paragraphs:
                    for run in para.runs:
                        run.bold = True
                row_idx += 1
            
            # Total G (Harga Satuan Pekerjaan)
            if row_idx < len(table.rows):
                table.rows[row_idx].cells[0].text = "Harga Satuan Pekerjaan (G = E + F)"
                for para in table.rows[row_idx].cells[0].paragraphs:
                    para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    for run in para.runs:
                        run.bold = True
                        run.font.size = Pt(10)
                table.rows[row_idx].cells[6].text = totals.get('G', '0')
                for para in table.rows[row_idx].cells[6].paragraphs:
                    for run in para.runs:
                        run.bold = True
                        run.font.size = Pt(10)
            
            # Column widths
            widths = [Mm(10), Mm(50), Mm(25), Mm(15), Mm(20), Mm(25), Mm(30)]
            for row in table.rows:
                for col_idx, cell in enumerate(row.cells):
                    cell.width = widths[col_idx]
            
            # Spacing between pekerjaan tables (no page break - more compact)
            if idx < len(sections) - 1:
                self.doc.add_paragraph()  # 1 line spacing
                self.doc.add_paragraph()  # 2 lines total
        
        # ========== SECTION 3: LEMBAR PENGESAHAN ==========
        # Add some spacing (not a new page, at bottom of last rincian)
        self.doc.add_paragraph()
        self.doc.add_paragraph()
        
        # Approval section title
        approval_title = self.doc.add_paragraph()
        approval_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        approval_run = approval_title.add_run('LEMBAR PENGESAHAN')
        approval_run.bold = True
        approval_run.font.size = Pt(12)
        
        self.doc.add_paragraph()
        
        # Approval table (2 columns for Pemilik Proyek and Konsultan Perencana)
        approval_table = self.doc.add_table(rows=5, cols=2)
        approval_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # Row 0: Headers
        approval_table.rows[0].cells[0].text = "Pemilik Proyek"
        approval_table.rows[0].cells[1].text = "Konsultan Perencana"
        for cell in approval_table.rows[0].cells:
            for para in cell.paragraphs:
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in para.runs:
                    run.bold = True
        
        # Row 1-3: Empty space for signature
        for i in range(1, 4):
            for cell in approval_table.rows[i].cells:
                cell.text = ""
        
        # Row 4: Name lines
        approval_table.rows[4].cells[0].text = "(_________________________)"
        approval_table.rows[4].cells[1].text = "(_________________________)"
        for cell in approval_table.rows[4].cells:
            for para in cell.paragraphs:
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Set column widths
        for row in approval_table.rows:
            row.cells[0].width = Mm(80)
            row.cells[1].width = Mm(80)
        
        return self._create_response('rincian_ahsp')
    
    def export_rekap(self, data: Dict[str, Any]) -> HttpResponse:
        """
        Export Rekap Laporan to Word document.
        
        Structure:
        1. Cover Page
        2. Table of Contents
        3. Grid Planned Section
        4. Grid Actual Section
        5. Kurva S Section
        
        Args:
            data: Export data containing:
                - project_info: Project metadata
                - planned_pages: Planned progress data
                - actual_pages: Actual progress data
                - kurva_s_data: Kurva S chart data
                - sections: TOC sections
                
        Returns:
            HttpResponse with .docx file
        """
        import time
        start = time.time()
        step_times = {}
        
        self.doc = Document()
        self._setup_page_layout('A3', 'landscape')
        
        project_info = data.get('project_info', {})
        
        # 1. Cover Page
        self._build_cover_page('rekap', project_info)
        self.doc.add_page_break()
        
        # 2. Table of Contents
        sections = data.get('sections', [])
        if sections:
            self._build_toc(sections)
            self.doc.add_page_break()
        
        # 3. Grid Planned Section
        planned_pages = data.get('planned_pages', [])
        if planned_pages:
            grid_planned_start = time.time()
            self._build_section_header('BAGIAN 1: GRID VIEW - RENCANA (PLANNED)')
            for i, page in enumerate(planned_pages):
                page_start = time.time()
                self._build_grid_table(page, mode='planned')
                print(f"[WordExporter] Grid Planned page {i+1}/{len(planned_pages)}: {time.time() - page_start:.2f}s")
            self.doc.add_page_break()
            step_times['grid_planned'] = time.time() - grid_planned_start
            print(f"[WordExporter] [TIME] Grid Planned ({len(planned_pages)} pages): {step_times['grid_planned']:.2f}s")
        
        # 4. Grid Actual Section
        actual_pages = data.get('actual_pages', [])
        if actual_pages:
            grid_actual_start = time.time()
            self._build_section_header('BAGIAN 2: GRID VIEW - REALISASI (ACTUAL)')
            for i, page in enumerate(actual_pages):
                page_start = time.time()
                self._build_grid_table(page, mode='actual')
                print(f"[WordExporter] Grid Actual page {i+1}/{len(actual_pages)}: {time.time() - page_start:.2f}s")
            self.doc.add_page_break()
            step_times['grid_actual'] = time.time() - grid_actual_start
            print(f"[WordExporter] [TIME] Grid Actual ({len(actual_pages)} pages): {step_times['grid_actual']:.2f}s")
        
        # NOTE: Gantt Chart dan Kurva S dihapus dari Word export
        # karena bukan format native Word (hanya embedded image).
        # User dapat download Gantt/Kurva S sebagai image terpisah dari UI.
        
        # Create response
        step_start = time.time()
        response = self._create_response('rekap_laporan')
        step_times['create_response'] = time.time() - step_start
        print(f"[WordExporter] [TIME] Create response (save doc): {step_times['create_response']:.2f}s")
        print(f"[WordExporter] [OK] Total export_rekap: {time.time() - start:.2f}s")
        
        return response
    
    def export_monthly(self, data: Dict[str, Any]) -> HttpResponse:
        """
        Export Laporan Bulanan to Word document.
        
        Structure:
        1. Cover Page
        2. Progress Pelaksanaan Page
        3. Kurva S Monthly (Landscape)
        4. Kurva S Portrait
        5. Signature Section
        
        Args:
            data: Export data for monthly report
                
        Returns:
            HttpResponse with .docx file
        """
        self.doc = Document()
        self._setup_page_layout('A4', 'portrait')
        
        project_info = data.get('project_info', {})
        month = data.get('month', 1)
        period_info = data.get('period', {})
        
        # 1. Cover Page
        self._build_cover_page('monthly', project_info, period_info)
        self.doc.add_page_break()
        
        # 2. Progress Pelaksanaan
        exec_summary = data.get('executive_summary', {})
        hierarchy_data = data.get('hierarchy_progress', [])
        self._build_progress_page(month, project_info, exec_summary, hierarchy_data, 'monthly')
        self.doc.add_page_break()
        
        # 3. Kurva S Monthly (switch to landscape)
        kurva_s_data = data.get('kurva_s_data', [])
        if kurva_s_data:
            self._add_section_break('landscape')
            self._build_section_header(f'RINGKASAN PROGRESS KURVA S (Bulan ke-{month})')
            self._build_kurva_s_section(kurva_s_data, data)
        
        # 4. Signature Section
        self._add_section_break('portrait')
        self._build_signature_section(project_info)
        
        return self._create_response(f'laporan_bulan_{month}')
    
    def export_weekly(self, data: Dict[str, Any]) -> HttpResponse:
        """
        Export Laporan Mingguan to Word document.
        
        Structure:
        1. Cover Page
        2. Weekly Progress Page
        
        Args:
            data: Export data for weekly report
                
        Returns:
            HttpResponse with .docx file
        """
        self.doc = Document()
        self._setup_page_layout('A4', 'portrait')
        
        project_info = data.get('project_info', {})
        week = data.get('week', 1)
        period_info = data.get('period', {})
        
        # 1. Cover Page
        self._build_cover_page('weekly', project_info, period_info)
        self.doc.add_page_break()
        
        # 2. Weekly Progress Page
        exec_summary = data.get('executive_summary', {})
        hierarchy_data = data.get('hierarchy_progress', [])
        self._build_progress_page(week, project_info, exec_summary, hierarchy_data, 'weekly')
        
        return self._create_response(f'laporan_minggu_{week}')
    
    # =========================================================================
    # PAGE LAYOUT SETUP
    # =========================================================================
    
    def _setup_page_layout(self, size: str = 'A4', orientation: str = 'portrait'):
        """
        Configure page size, orientation, and margins.
        
        Args:
            size: 'A4' or 'A3'
            orientation: 'portrait' or 'landscape'
        """
        section = self.doc.sections[0]
        
        # A4: 210mm x 297mm (Portrait), A3: 297mm x 420mm (Portrait)
        if size.upper() == 'A3':
            base_width = Mm(297)
            base_height = Mm(420)
        else:  # A4
            base_width = Mm(210)
            base_height = Mm(297)
        
        # Apply orientation
        if orientation == 'landscape':
            section.orientation = WD_ORIENT.LANDSCAPE
            section.page_width = base_height  # Swap for landscape
            section.page_height = base_width
        else:  # portrait
            section.orientation = WD_ORIENT.PORTRAIT
            section.page_width = base_width
            section.page_height = base_height
        
        # Margins from config
        section.top_margin = Mm(self.config.margin_top)
        section.bottom_margin = Mm(self.config.margin_bottom)
        section.left_margin = Mm(self.config.margin_left)
        section.right_margin = Mm(self.config.margin_right)
    
    def _add_section_break(self, orientation: str = 'portrait'):
        """
        Add section break with new orientation.
        
        Args:
            orientation: 'portrait' or 'landscape'
        """
        new_section = self.doc.add_section()
        
        if orientation == 'landscape':
            new_section.orientation = WD_ORIENT.LANDSCAPE
            new_section.page_width = Mm(420 if self.config.page_size == 'A3' else 297)
            new_section.page_height = Mm(297 if self.config.page_size == 'A3' else 210)
        else:
            new_section.orientation = WD_ORIENT.PORTRAIT
            new_section.page_width = Mm(297 if self.config.page_size == 'A3' else 210)
            new_section.page_height = Mm(420 if self.config.page_size == 'A3' else 297)
        
        # Copy margins
        new_section.top_margin = Mm(self.config.margin_top)
        new_section.bottom_margin = Mm(self.config.margin_bottom)
        new_section.left_margin = Mm(self.config.margin_left)
        new_section.right_margin = Mm(self.config.margin_right)
    
    def _build_signature_section(self, project_info: Dict[str, Any] = None):
        """
        Build signature section with signature boxes.
        Uses config.signature_config for signature data.
        
        Args:
            project_info: Optional project info dict (for backward compatibility)
        """
        # Add title
        title_para = self.doc.add_paragraph()
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title_para.add_run('LEMBAR PENGESAHAN')
        title_run.bold = True
        title_run.font.size = Pt(14)
        
        self.doc.add_paragraph()  # Spacing
        
        # Get signature data from config
        sig_config = self.config.signature_config
        signatures = sig_config.signatures if sig_config and sig_config.enabled else []
        
        if not signatures:
            # Fallback to default signatures if none configured
            signatures = [
                {'label': 'Pemilik Proyek', 'name': '', 'position': ''},
                {'label': 'Konsultan Perencana', 'name': '', 'position': ''},
            ]
        
        # Create 2-column table for signatures (side-by-side)
        num_cols = min(len(signatures), 3)
        table = self.doc.add_table(rows=4, cols=num_cols)
        
        for col_idx, sig in enumerate(signatures[:num_cols]):
            # Row 0: Label
            cell0 = table.rows[0].cells[col_idx]
            cell0.text = sig.get('label', '')
            for para in cell0.paragraphs:
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in para.runs:
                    run.bold = True
            
            # Row 1: Blank space for actual signature
            cell1 = table.rows[1].cells[col_idx]
            cell1.text = ''
            cell1.paragraphs[0].add_run('\n\n\n')  # Signature space
            
            # Row 2: Name
            cell2 = table.rows[2].cells[col_idx]
            name = sig.get('name', '')
            cell2.text = name if name else '________________________'
            for para in cell2.paragraphs:
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Row 3: Position
            cell3 = table.rows[3].cells[col_idx]
            position = sig.get('position', '')
            cell3.text = position if position else ''
            for para in cell3.paragraphs:
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    def _build_pengesahan_word_table(self, table_data: Dict[str, Any], col_widths: List = None):
        """
        Build dedicated 3-column table for REKAPITULASI RENCANA ANGGARAN BIAYA.
        
        This is separate from the generic table building to ensure pengesahan page
        only has 3 columns: No, Uraian Klasifikasi, Jumlah Harga
        
        Args:
            table_data: Dict with 'headers' and 'rows'
            col_widths: List of column widths in mm [15, 105, 70]
        """
        headers = table_data.get('headers', ['No', 'Uraian Pekerjaan', 'Jumlah Harga (Rp)'])
        rows = table_data.get('rows', [])
        
        if not rows:
            para = self.doc.add_paragraph("Tidak ada data")
            return
        
        # Default col widths if not provided
        if not col_widths:
            col_widths = [15, 105, 70]
        
        # Create table with exactly 3 columns
        num_cols = 3
        num_rows = len(rows) + 1  # +1 for header
        
        table = self.doc.add_table(rows=num_rows, cols=num_cols)
        table.style = 'Table Grid'
        
        # Header row
        for col_idx, header_text in enumerate(headers[:num_cols]):
            cell = table.rows[0].cells[col_idx]
            cell.text = str(header_text)
            # Style header cell (bold)
            for para in cell.paragraphs:
                for run in para.runs:
                    run.bold = True
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Data rows
        for row_idx, row_data in enumerate(rows):
            table_row = table.rows[row_idx + 1]
            
            # Get exactly 3 values
            if len(row_data) >= 3:
                values = [row_data[0], row_data[1], row_data[2]]
            elif len(row_data) == 2:
                values = ['', row_data[0], row_data[1]]
            else:
                values = [row_data[0] if row_data else '', '', '']
            
            for col_idx, value in enumerate(values):
                cell = table_row.cells[col_idx]
                cell.text = str(value) if value else ''
                
                # Column alignment
                for para in cell.paragraphs:
                    if col_idx == 0:  # No column - center
                        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    elif col_idx == 1:  # Uraian - left
                        para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    else:  # Jumlah - right
                        para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        # Set column widths
        for row in table.rows:
            for col_idx, cell in enumerate(row.cells):
                if col_idx < len(col_widths):
                    cell.width = Mm(col_widths[col_idx])
    
    # =========================================================================
    # COVER PAGE
    # =========================================================================
    
    def _build_cover_page(self, report_type: str, project_info: Dict[str, Any], 
                          period_info: Dict[str, Any] = None):
        """
        Build cover page with title and project identity.
        
        Args:
            report_type: 'rekap', 'monthly', or 'weekly'
            project_info: Project information dict
            period_info: Period info for monthly/weekly reports
        """
        # Title based on report type
        if report_type == 'rekap':
            title = 'REKAP LAPORAN JADWAL PEKERJAAN'
        elif report_type == 'monthly':
            month = period_info.get('month', 1) if period_info else 1
            title = f'LAPORAN BULAN KE-{month}'
        else:  # weekly
            week = period_info.get('week', 1) if period_info else 1
            title = f'LAPORAN MINGGU KE-{week}'
        
        # Add spacing at top
        for _ in range(3):
            self.doc.add_paragraph()
        
        # Main title
        title_para = self.doc.add_paragraph()
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title_para.add_run(title)
        title_run.bold = True
        title_run.font.size = Pt(24)
        title_run.font.color.rgb = RGBColor.from_string(UTS.PRIMARY_LIGHT[1:])
        
        # Subtitle - Project name
        project_name = project_info.get('nama_project', self.config.project_name)
        if project_name:
            subtitle_para = self.doc.add_paragraph()
            subtitle_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            subtitle_run = subtitle_para.add_run(project_name)
            subtitle_run.bold = True
            subtitle_run.font.size = Pt(18)
        
        # Spacing
        self.doc.add_paragraph()
        self.doc.add_paragraph()
        
        # Project identity table
        identity_rows = build_identity_rows(self.config)
        if identity_rows:
            table = self.doc.add_table(rows=len(identity_rows), cols=3)
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            
            for i, row_data in enumerate(identity_rows):
                row = table.rows[i]
                for j, cell_text in enumerate(row_data):
                    cell = row.cells[j]
                    cell.text = str(cell_text)
                    # Style
                    for para in cell.paragraphs:
                        for run in para.runs:
                            run.font.size = Pt(11)
    
    # =========================================================================
    # TABLE OF CONTENTS
    # =========================================================================
    
    def _build_toc(self, sections: List[str]):
        """
        Build table of contents.
        
        Args:
            sections: List of section titles
        """
        # TOC Title
        toc_title = self.doc.add_paragraph()
        toc_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        toc_run = toc_title.add_run('DAFTAR ISI')
        toc_run.bold = True
        toc_run.font.size = Pt(16)
        
        self.doc.add_paragraph()
        
        # Section list
        for idx, section_name in enumerate(sections, 1):
            para = self.doc.add_paragraph()
            para.add_run(f'{idx}. {section_name}')
    
    # =========================================================================
    # SECTION HEADERS
    # =========================================================================
    
    def _build_section_header(self, title: str):
        """
        Build section header with styling.
        
        Args:
            title: Section title text
        """
        para = self.doc.add_paragraph()
        run = para.add_run(title)
        run.bold = True
        run.font.size = Pt(14)
        run.font.color.rgb = RGBColor.from_string(UTS.PRIMARY_LIGHT[1:])
        
        self.doc.add_paragraph()  # Spacing
    
    # =========================================================================
    # GRID TABLE
    # =========================================================================
    
    def _build_grid_table(self, page_data: Dict[str, Any], mode: str = 'planned'):
        """
        Build grid table with hierarchy styling.
        
        Args:
            page_data: Page data containing table_data, headers
            mode: 'planned' or 'actual'
        """
        table_data = page_data.get('table_data', {})
        headers = table_data.get('headers', [])
        rows = table_data.get('rows', [])
        hierarchy_levels = page_data.get('hierarchy_levels', {})
        
        if not headers or not rows:
            return
        
        num_cols = len(headers)
        num_rows = len(rows) + 1  # +1 for header
        
        # Create table
        table = self.doc.add_table(rows=num_rows, cols=num_cols)
        table.style = 'Table Grid'
        
        # Header row
        header_row = table.rows[0]
        for col_idx, header_text in enumerate(headers):
            cell = header_row.cells[col_idx]
            cell.text = str(header_text)
            self._style_header_cell(cell)
        
        # Enable header repeat
        self._enable_header_repeat(table)
        
        # Data rows
        for row_idx, row_data in enumerate(rows):
            table_row = table.rows[row_idx + 1]
            level = hierarchy_levels.get(row_idx, 3)
            row_type = self._get_row_type_from_level(level)
            
            for col_idx, cell_value in enumerate(row_data):
                cell = table_row.cells[col_idx]
                
                # Skip 0% values to reduce clutter and file size
                display_value = ''
                if cell_value:
                    str_val = str(cell_value).strip()
                    # Skip if value is 0, 0%, 0.0, 0.0%, etc.
                    # Include both dot (.) and comma (,) decimal formats for Indonesian locale
                    zero_values = (
                        '0', '0%',
                        '0.0', '0.0%', '0.00', '0.00%', '0.000', '0.000%',
                        '0,0', '0,0%', '0,00', '0,00%', '0,000', '0,000%',
                    )
                    if str_val not in zero_values:
                        display_value = str_val
                cell.text = display_value
                
                # Apply hierarchy styling
                self._apply_hierarchy_style(cell, row_type, col_idx)
        
        # Set column widths
        self._set_grid_column_widths(table, num_cols)
        
        self.doc.add_paragraph()  # Spacing after table
    
    def _style_header_cell(self, cell):
        """Apply header cell styling."""
        # Background color
        shading = OxmlElement('w:shd')
        shading.set(qn('w:fill'), UTS.PRIMARY_LIGHT[1:])
        cell._tc.get_or_add_tcPr().append(shading)
        
        # Text styling
        for para in cell.paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in para.runs:
                run.bold = True
                run.font.size = Pt(7)  # 7pt for headers
                run.font.name = 'Arial'
                run.font.color.rgb = RGBColor(255, 255, 255)
    
    def _apply_hierarchy_style(self, cell, row_type: str, col_idx: int):
        """
        Apply hierarchy-based styling to cell.
        
        Args:
            cell: Table cell
            row_type: 'klasifikasi', 'sub_klasifikasi', or 'pekerjaan'
            col_idx: Column index
        """
        # Get background color and font settings
        # Font: Arial 7pt for data rows (per user request)
        if row_type == 'klasifikasi':
            bg_color = UTS.KLASIFIKASI_BG[1:]
            bold = True
            font_size = 8  # Slightly larger for main category
        elif row_type == 'sub_klasifikasi':
            bg_color = UTS.SUB_KLASIFIKASI_BG[1:]
            bold = True
            font_size = 7
        else:  # pekerjaan
            bg_color = 'FFFFFF'
            bold = False
            font_size = 7  # 7pt for data rows
        
        # Apply background
        shading = OxmlElement('w:shd')
        shading.set(qn('w:fill'), bg_color)
        cell._tc.get_or_add_tcPr().append(shading)
        
        # Text styling
        for para in cell.paragraphs:
            if col_idx == 0:  # Uraian column - left align
                para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            else:  # Other columns - center
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            for run in para.runs:
                run.bold = bold
                run.font.size = Pt(font_size)
                run.font.name = 'Arial'  # Base font Arial
    
    def _enable_header_repeat(self, table):
        """Enable table header row repeat on page breaks."""
        tbl = table._tbl
        for row in table.rows[:1]:
            tr = row._tr
            trPr = tr.get_or_add_trPr()
            tblHeader = OxmlElement('w:tblHeader')
            trPr.append(tblHeader)
    
    def _set_grid_column_widths(self, table, num_cols: int):
        """
        Set appropriate column widths for grid table.
        
        Args:
            table: Word table object
            num_cols: Number of columns
        """
        # Static columns: Uraian, Volume, Satuan = 3 cols
        # Rest are week columns
        static_cols = 3
        week_cols = max(0, num_cols - static_cols)
        
        # Calculate widths
        uraian_width = Mm(JadwalExportLayout.COL_URAIAN)
        volume_width = Mm(JadwalExportLayout.COL_VOLUME)
        satuan_width = Mm(JadwalExportLayout.COL_SATUAN)
        
        # Week width - dynamic based on count
        min_week_mm = 4.2  # ~12pt
        max_week_mm = 15.9  # ~45pt
        
        if week_cols > 0:
            # Calculate available width (A3 landscape - margins - static cols)
            available_mm = 396 - JadwalExportLayout.COL_URAIAN - JadwalExportLayout.COL_VOLUME - JadwalExportLayout.COL_SATUAN
            week_width_mm = available_mm / week_cols
            week_width_mm = max(min_week_mm, min(max_week_mm, week_width_mm))
            week_width = Mm(week_width_mm)
        else:
            week_width = Mm(15)
        
        # Apply widths
        for row in table.rows:
            for col_idx, cell in enumerate(row.cells):
                if col_idx == 0:
                    cell.width = uraian_width
                elif col_idx == 1:
                    cell.width = volume_width
                elif col_idx == 2:
                    cell.width = satuan_width
                else:
                    cell.width = week_width
    
    def _get_row_type_from_level(self, level: int) -> str:
        """Convert hierarchy level to row type string."""
        if level == 1:
            return 'klasifikasi'
        elif level == 2:
            return 'sub_klasifikasi'
        else:
            return 'pekerjaan'
    
    # =========================================================================
    # ATTACHMENT IMAGE EMBEDDING
    # =========================================================================
    
    def _embed_attachment_image(self, attachment: Dict[str, Any]):
        """
        Embed attachment image into document.
        
        Used for Gantt charts and Kurva S rendered by frontend.
        
        Args:
            attachment: Dict with 'title', 'bytes' (base64), 'format'
        """
        import base64
        
        title = attachment.get('title', 'Chart')
        img_bytes = attachment.get('bytes', '')
        img_format = attachment.get('format', 'png')
        
        if not img_bytes:
            # Add placeholder text if no image
            para = self.doc.add_paragraph()
            para.add_run(f'[{title} - Image not available]')
            return
        
        try:
            image_data = None
            
            # Case 1: Already raw image bytes (PNG starts with \x89PNG)
            if isinstance(img_bytes, bytes):
                # Check if it's raw PNG (starts with \x89PNG) or JPEG (\xFF\xD8)
                if img_bytes[:4] == b'\x89PNG' or img_bytes[:2] == b'\xff\xd8':
                    # Already raw image bytes, use directly
                    image_data = img_bytes
                else:
                    # Might be base64 encoded as bytes, decode to string first
                    try:
                        img_str = img_bytes.decode('ascii')
                        # Remove data URL prefix if present
                        if img_str.startswith('data:'):
                            if ',' in img_str:
                                img_str = img_str.split(',', 1)[1]
                        image_data = base64.b64decode(img_str)
                    except (UnicodeDecodeError, Exception):
                        # Not valid base64, try using as-is
                        image_data = img_bytes
            
            # Case 2: String (base64 or data URL)
            elif isinstance(img_bytes, str):
                # Remove data URL prefix if present
                if img_bytes.startswith('data:'):
                    if ',' in img_bytes:
                        img_bytes = img_bytes.split(',', 1)[1]
                image_data = base64.b64decode(img_bytes)
            
            if not image_data:
                raise ValueError("Could not process image data")
            
            image_stream = BytesIO(image_data)
            
            # Add title
            title_para = self.doc.add_paragraph()
            title_run = title_para.add_run(title)
            title_run.bold = True
            title_run.font.size = Pt(10)
            title_run.font.name = 'Arial'
            
            # Add image - fit to page width (A3 landscape ~ 39cm usable width)
            # Use 35cm width to leave margins
            self.doc.add_picture(image_stream, width=Cm(35))
            
            # Add spacing after image
            self.doc.add_paragraph()
            
        except Exception as e:
            # If image embedding fails, add error message
            para = self.doc.add_paragraph()
            para.add_run(f'[{title} - Error embedding image: {str(e)}]')
    
    # =========================================================================
    # KURVA S SECTION
    # =========================================================================
    
    def _build_kurva_s_section(self, kurva_s_data: List[Dict], data: Dict[str, Any]):
        """
        Build Kurva S section (table-based visualization).
        
        Args:
            kurva_s_data: Chart data points
            data: Full export data
        """
        # For now, create a simple summary table
        # TODO: Add chart image or table-based visualization
        
        if not kurva_s_data:
            return
        
        para = self.doc.add_paragraph()
        para.add_run('Kurva S data visualization will be added here.')
        
        # Create simple data table
        num_weeks = len(kurva_s_data)
        table = self.doc.add_table(rows=3, cols=min(num_weeks + 1, 20))
        table.style = 'Table Grid'
        
        # Headers
        table.rows[0].cells[0].text = 'Minggu'
        table.rows[1].cells[0].text = 'Rencana (%)'
        table.rows[2].cells[0].text = 'Realisasi (%)'
        
        # Data
        for idx, week_data in enumerate(kurva_s_data[:19]):
            col_idx = idx + 1
            if col_idx < len(table.rows[0].cells):
                table.rows[0].cells[col_idx].text = str(week_data.get('week', idx + 1))
                table.rows[1].cells[col_idx].text = f"{week_data.get('planned', 0):.1f}"
                table.rows[2].cells[col_idx].text = f"{week_data.get('actual', 0):.1f}"
    
    # =========================================================================
    # PROGRESS PAGE
    # =========================================================================
    
    def _build_progress_page(self, period: int, project_info: Dict[str, Any],
                             summary: Dict[str, Any], hierarchy_data: List[Dict],
                             mode: str = 'monthly'):
        """
        Build progress page for monthly/weekly reports.
        
        Args:
            period: Month or week number
            project_info: Project information
            summary: Executive summary data
            hierarchy_data: Hierarchy progress data
            mode: 'monthly' or 'weekly'
        """
        period_label = 'Bulan' if mode == 'monthly' else 'Minggu'
        
        # Section title
        self._build_section_header(f'PROGRESS PELAKSANAAN - {period_label} ke-{period}')
        
        # Summary table
        if summary:
            self._build_summary_table(summary)
            self.doc.add_paragraph()
        
        # Hierarchy progress table
        if hierarchy_data:
            self._build_hierarchy_progress_table(hierarchy_data)
    
    def _build_summary_table(self, summary: Dict[str, Any]):
        """Build executive summary table."""
        table = self.doc.add_table(rows=4, cols=2)
        table.style = 'Table Grid'
        
        data = [
            ('Progress Rencana', f"{summary.get('planned_progress', 0):.2f}%"),
            ('Progress Realisasi', f"{summary.get('actual_progress', 0):.2f}%"),
            ('Deviasi', f"{summary.get('deviation', 0):.2f}%"),
            ('Status', summary.get('status', '-')),
        ]
        
        for idx, (label, value) in enumerate(data):
            table.rows[idx].cells[0].text = label
            table.rows[idx].cells[1].text = str(value)
    
    def _build_hierarchy_progress_table(self, hierarchy_data: List[Dict]):
        """Build hierarchy progress detail table."""
        if not hierarchy_data:
            return
        
        # Create table with hierarchy data
        table = self.doc.add_table(rows=len(hierarchy_data) + 1, cols=4)
        table.style = 'Table Grid'
        
        # Headers
        headers = ['Uraian', 'Rencana (%)', 'Realisasi (%)', 'Deviasi (%)']
        for idx, header in enumerate(headers):
            table.rows[0].cells[idx].text = header
            self._style_header_cell(table.rows[0].cells[idx])
        
        # Data rows
        for row_idx, item in enumerate(hierarchy_data):
            row = table.rows[row_idx + 1]
            row.cells[0].text = item.get('name', '')
            row.cells[1].text = f"{item.get('planned', 0):.2f}"
            row.cells[2].text = f"{item.get('actual', 0):.2f}"
            row.cells[3].text = f"{item.get('deviation', 0):.2f}"

    # =========================================================================
    # RESPONSE CREATION
    # =========================================================================
    
    def _create_response(self, filename: str) -> HttpResponse:
        """
        Create HTTP response with Word document.
        
        Args:
            filename: Base filename (without extension)
            
        Returns:
            HttpResponse with .docx attachment
        """
        buffer = BytesIO()
        self.doc.save(buffer)
        buffer.seek(0)
        
        response = HttpResponse(
            buffer.getvalue(),
            content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        response['Content-Disposition'] = f'attachment; filename="{filename}.docx"'
        
        return response
